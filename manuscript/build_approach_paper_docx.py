from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(
    "/Users/amoghn/Desktop/IEEM Dr. Zhang/Approach Paper/"
    "Approach Paper for MISO vs SISO Draft 1.docx"
)
OUTPUT = Path(__file__).resolve().parent / "Approach Paper for MISO vs SISO Draft 1.docx"


NAVY = "17365D"
BLUE = "285F8F"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FB"
AMBER = "9C6500"
LIGHT_AMBER = "FFF4D6"
GRAY = "5B6573"
LIGHT_GRAY = "F2F3F5"
MID_GRAY = "D8DDE3"
WHITE = "FFFFFF"
BLACK = "111111"
RED = "A61B1B"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_borders(table, color=MID_GRAY, size="4") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_cell_border(cell, *, left=None, right=None, top=None, bottom=None) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in (
        ("left", left),
        ("right", right),
        ("top", top),
        ("bottom", bottom),
    ):
        if edge_data is None:
            continue
        tag = tc_borders.find(qn(f"w:{edge_name}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge_name}")
            tc_borders.append(tag)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                tag.set(qn(f"w:{key}"), str(edge_data[key]))


def set_run_font(run, name: str, size: float | None = None, color: str | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(fld_char_end)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 21, NAVY, 0, 12),
        ("Subtitle", 12, GRAY, 0, 18),
        ("Heading 1", 15, NAVY, 16, 8),
        ("Heading 2", 12.5, BLUE, 13, 5),
        ("Heading 3", 11.5, NAVY, 10, 3),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Equation" not in styles:
        equation_style = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation_style = styles["Equation"]
    equation_style.font.name = "Cambria Math"
    equation_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation_style.font.size = Pt(10.5)
    equation_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_style.paragraph_format.space_before = Pt(5)
    equation_style.paragraph_format.space_after = Pt(5)
    equation_style.paragraph_format.keep_together = True

    if "Caption Custom" not in styles:
        cap = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Custom"]
    cap.font.name = "Arial"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor.from_string(GRAY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_with_next = False

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    ref.font.size = Pt(10)
    ref.paragraph_format.left_indent = Inches(0.28)
    ref.paragraph_format.first_line_indent = Inches(-0.28)
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.line_spacing = 1.05

    if "Code" not in styles:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.2)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = True

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("MISO versus SISO approach paper • working draft")
        set_run_font(r, "Arial", 8, GRAY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("Draft 1  |  ")
        set_run_font(r1, "Arial", 8, GRAY)
        add_field(fp, "PAGE", "1")
        for run in fp.runs:
            set_run_font(run, "Arial", 8, GRAY)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = (
        "A Practical Two-Input Transfer Function Framework for Cerebral Hemodynamics"
    )
    props.subject = "Working draft of the MISO versus SISO approach paper"
    props.keywords = (
        "dynamic cerebral autoregulation; transfer function analysis; MISO; "
        "SISO; end-tidal carbon dioxide; cerebral blood flow velocity"
    )


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align=None,
    bold_prefix: str | None = None,
    italic: bool = False,
) -> object:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix) :])
        r2.italic = italic
    else:
        r = p.add_run(text)
        r.italic = italic
    return p


def add_bullets(doc: Document, items: Iterable[str], *, level=0) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.28 + level * 0.22)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("• ")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(f"{index}. ")
        p.add_run(item)


def add_equation(doc: Document, text: str, number: int | str | None = None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Equation"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.4), WD_TAB_ALIGNMENT.CENTER
    )
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.8), WD_TAB_ALIGNMENT.RIGHT
    )
    r = p.add_run("\t" + text)
    set_run_font(r, "Cambria Math", 10.5, BLACK)
    if number is not None:
        rn = p.add_run(f"\t({number})")
        set_run_font(rn, "Times New Roman", 10, GRAY)


def add_callout(
    doc: Document,
    title: str,
    body: str,
    *,
    kind: str = "note",
) -> None:
    fill = LIGHT_AMBER if kind == "note" else LIGHT_BLUE
    accent = AMBER if kind == "note" else BLUE
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.8)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "16", "space": "0", "color": accent},
        top={"val": "single", "sz": "4", "space": "0", "color": accent},
        bottom={"val": "single", "sz": "4", "space": "0", "color": accent},
        right={"val": "single", "sz": "4", "space": "0", "color": accent},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_together = True
    rt = p.add_run(title)
    rt.bold = True
    set_run_font(rt, "Arial", 9.5, accent)
    rb = p.add_run("\n" + body)
    set_run_font(rb, "Times New Roman", 10.2, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure_placeholder(
    doc: Document,
    number: str,
    title: str,
    content: str,
    purpose: str,
    *,
    height_in: float = 1.25,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.8)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "10", "space": "0", "color": BLUE},
        top={"val": "single", "sz": "10", "space": "0", "color": BLUE},
        right={"val": "single", "sz": "10", "space": "0", "color": BLUE},
        bottom={"val": "single", "sz": "10", "space": "0", "color": BLUE},
    )
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_in * 1440)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)
    prevent_row_split(table.rows[0])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    r = p.add_run(f"FIGURE PLACEHOLDER — Figure {number}\n")
    r.bold = True
    set_run_font(r, "Arial", 10.5, NAVY)
    r2 = p.add_run(title + "\n")
    r2.bold = True
    set_run_font(r2, "Arial", 10, BLUE)
    r3 = p.add_run("What to show: " + content + "\n")
    set_run_font(r3, "Times New Roman", 9.5, BLACK)
    r4 = p.add_run("Why it matters: " + purpose)
    r4.italic = True
    set_run_font(r4, "Times New Roman", 9.3, GRAY)
    cap = doc.add_paragraph(style="Caption Custom")
    rc = cap.add_run(f"Figure {number}. {title}. ")
    rc.bold = True
    cap.add_run(
        "Replace this placeholder with the final figure and revise the caption "
        "after the corresponding analysis is complete."
    )


def add_code_block(
    doc: Document,
    title: str,
    code: str,
    note: str,
    *,
    placeholder: bool = False,
) -> None:
    if doc.paragraphs:
        # Keep a normal body paragraph with the table so a code block does not
        # become the first item on a page; LibreOffice otherwise suppresses the
        # visible header/footer margin when rendering some table-led pages.
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
    table = doc.add_table(rows=2, cols=1)
    set_table_fixed(table)
    set_table_borders(table, color="C5CBD2", size="4")
    set_cell_width(table.cell(0, 0), 6.8)
    set_cell_width(table.cell(1, 0), 6.8)
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE if not placeholder else LIGHT_AMBER)
    set_cell_shading(table.cell(1, 0), LIGHT_GRAY)
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    hp = table.cell(0, 0).paragraphs[0]
    hp.paragraph_format.keep_with_next = True
    hr = hp.add_run(("CODE PLACEHOLDER — " if placeholder else "CODE EXCERPT — ") + title)
    hr.bold = True
    set_run_font(hr, "Arial", 9.2, AMBER if placeholder else NAVY)
    cp = table.cell(1, 0).paragraphs[0]
    cp.style = doc.styles["Code"]
    for index, line in enumerate(code.splitlines()):
        if index:
            cp.add_run().add_break()
        rr = cp.add_run(line)
        set_run_font(rr, "Consolas", 8.2, BLACK)
    np = doc.add_paragraph(style="Caption Custom")
    nr = np.add_run(note)
    nr.italic = True


def add_table(
    doc: Document,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[float],
    note: str | None = None,
    font_size: float = 8.6,
) -> None:
    cap = doc.add_paragraph(style="Caption Custom")
    cap.paragraph_format.keep_with_next = True
    rr = cap.add_run(title)
    rr.bold = True
    set_run_font(rr, "Arial", 9.2, NAVY)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(table)
    set_table_borders(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for j, value in enumerate(headers):
        cell = header.cells[j]
        set_cell_width(cell, widths[j])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=85, start=65, bottom=85, end=65)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(value))
        r.bold = True
        set_run_font(r, "Arial", font_size, WHITE)
    for i, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            set_cell_width(cell, widths[j])
            set_cell_margins(cell, top=65, start=60, bottom=65, end=60)
            if i % 2 == 1:
                set_cell_shading(cell, "F8F9FA")
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, "Times New Roman", font_size, BLACK)
    if note:
        np = doc.add_paragraph(style="Caption Custom")
        np.add_run(note)


def add_section(doc: Document, title: str, level: int = 1) -> None:
    doc.add_heading(title, level=level)


def build_document() -> None:
    doc = Document(SOURCE)
    clear_document_body(doc)
    configure_document(doc)

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------
    p = doc.add_paragraph(style="Title")
    p.add_run(
        "A Practical Two-Input Transfer Function Framework for Cerebral "
        "Hemodynamics"
    )
    sp = doc.add_paragraph(style="Subtitle")
    sp.add_run(
        "Separating mean arterial pressure and end-tidal carbon dioxide "
        "effects: theory, validation plan, and an empirical demonstration "
        "in normal-control participants"
    )
    add_paragraph(
        doc,
        "[Author list to be confirmed]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "[Affiliations to be confirmed]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
    add_paragraph(
        doc,
        "Working Draft 1 — July 23, 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_callout(
        doc,
        "Purpose of this draft",
        "This is the approach paper only. The biological-marker paper will be "
        "separate. This manuscript uses the normal-control (NC) recordings as "
        "the empirical demonstration and does not make an NC-versus-MCI claim. "
        "Completed NC findings are written as results. Simulation, surrogate, "
        "delay, and regularization analyses that have not yet been run are "
        "marked clearly as pending. These markers should remain visible until "
        "the corresponding result is available.",
        kind="info",
    )
    add_callout(
        doc,
        "Title decision before submission",
        "If the simulation and robustness sections are completed as planned, "
        "the title can use the word “validation.” If those analyses are not "
        "completed, keep the more cautious word “framework.”",
    )
    doc.add_page_break()

    # ------------------------------------------------------------------
    # Abstract and front matter
    # ------------------------------------------------------------------
    add_section(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Background: Transfer function analysis of dynamic cerebral "
        "autoregulation is commonly performed with mean arterial pressure "
        "(MAP) as the only input and cerebral blood flow velocity (CBFV) as "
        "the output. CBFV also responds to arterial carbon dioxide. When MAP "
        "and end-tidal carbon dioxide (PₑₜCO₂) vary together, a MAP-only "
        "single-input, single-output (SISO) estimate can contain information "
        "that is shared with PₑₜCO₂.",
        bold_prefix="Background:",
    )
    add_paragraph(
        doc,
        "Objective: To develop and validate a transparent two-input, "
        "single-output (MISO) frequency-domain framework that estimates the "
        "conditional MAP-to-CBFV and PₑₜCO₂-to-CBFV transfer functions, and "
        "to define the conditions under which the added input improves or "
        "destabilizes estimation.",
        bold_prefix="Objective:",
    )
    add_paragraph(
        doc,
        "Methods: The two-input estimator was derived from the 2 × 2 input "
        "spectral matrix. A known-truth Monte Carlo study was designed to vary "
        "input coherence, PₑₜCO₂ contribution, output noise, recording "
        "duration, PₑₜCO₂ delay, model misspecification, and regularization. "
        "The empirical demonstration used resting MAP, PₑₜCO₂, and CBFV "
        "recordings from 23 NC participants. Signals were resampled to 4 Hz. "
        "Spectra were estimated with 128-second Hann windows, 50% overlap, and "
        "a three-bin smoothing kernel. MISO and SISO gain were compared within "
        "participant across four prespecified frequency bands.",
        bold_prefix="Methods:",
    )
    add_paragraph(
        doc,
        "Results: [INSERT SIMULATION RESULTS AFTER THE KNOWN-TRUTH ANALYSIS IS "
        "RUN.] In the current NC demonstration, MISO MAP gain was lower than "
        "SISO MAP gain in the very-very-low-frequency band (mean difference "
        "−0.240 percentage points of baseline CBFV per mmHg; 95% CI, −0.384 "
        "to −0.096; paired d_z = −0.72; median MISO/SISO ratio, 0.837). A small "
        "increase was present in the low-frequency band (difference 0.065; "
        "95% CI, 0.010 to 0.120; d_z = 0.51; median ratio, 1.028). The very-low- "
        "and high-frequency MAP differences were not supported. PₑₜCO₂-only "
        "SISO gain was substantially larger than conditional MISO PₑₜCO₂ "
        "gain at low and high frequencies, but these estimates require caution "
        "because PₑₜCO₂ power and ordinary coherence were low in parts of that "
        "range.",
        bold_prefix="Results:",
    )
    add_paragraph(
        doc,
        "Conclusions: A two-input transfer function is best understood as a "
        "conditional estimator, not as an automatic upgrade over SISO. The NC "
        "data show that adding PₑₜCO₂ can materially change the estimated "
        "slow MAP-to-CBFV relationship. The final methodological claim should "
        "depend on known-truth simulation, scale-invariant conditioning, delay "
        "sensitivity, surrogate-input testing, and prespecified regularization "
        "analyses.",
        bold_prefix="Conclusions:",
    )
    add_paragraph(
        doc,
        "Keywords: dynamic cerebral autoregulation; transfer function analysis; "
        "multiple input; end-tidal carbon dioxide; cerebral blood flow velocity; "
        "spectral analysis; system identification",
        bold_prefix="Keywords:",
    )
    add_callout(
        doc,
        "Abstract completion note",
        "The abstract is intentionally honest about the unfinished validation "
        "analyses. Replace the bracketed simulation sentence only after the "
        "final simulation code and sensitivity analyses are frozen.",
    )

    add_section(doc, "Plain-language overview", 1)
    add_paragraph(
        doc,
        "A standard cerebral autoregulation analysis asks how changes in blood "
        "pressure are followed by changes in blood-flow velocity in the brain. "
        "That question is useful, but blood pressure is not the only signal "
        "changing at the same time. Carbon dioxide is a strong dilator of "
        "cerebral vessels, and spontaneous changes in breathing can alter both "
        "end-tidal carbon dioxide and cerebral blood flow. If pressure and "
        "carbon dioxide also move together, a pressure-only model cannot always "
        "tell which input is responsible for the shared part of the response.",
    )
    add_paragraph(
        doc,
        "The two-input model estimates pressure and carbon dioxide at the same "
        "time. This can reduce omitted-input error, but it creates a new "
        "problem: the model can become unstable when carbon dioxide barely "
        "changes or when the two inputs are too similar. This paper therefore "
        "does not ask whether MISO is always better. It asks a more useful "
        "question: under what data conditions is the MISO answer accurate, "
        "stable, and worth interpreting?",
    )
    add_callout(
        doc,
        "Journal note",
        "Keep this plain-language section if the target journal requests a "
        "lay summary. Otherwise, its clearest sentences can be moved into the "
        "Introduction and the section can be removed.",
    )

    # ------------------------------------------------------------------
    # Introduction
    # ------------------------------------------------------------------
    add_section(doc, "1. Introduction", 1)
    add_section(doc, "1.1 Dynamic cerebral autoregulation and transfer function analysis", 2)
    add_paragraph(
        doc,
        "The brain requires a continuous supply of oxygen and metabolic "
        "substrates. Cerebral blood flow is therefore regulated by several "
        "overlapping mechanisms, including changes in cerebrovascular "
        "resistance, arterial pressure, carbon dioxide, oxygen, autonomic "
        "activity, cardiac output, and local neural activity. Dynamic cerebral "
        "autoregulation (dCA) describes the time-dependent pressure–flow "
        "relationship over seconds to minutes. It is commonly studied by "
        "recording arterial pressure and cerebral blood flow velocity (CBFV) "
        "and asking how oscillations in pressure are transmitted to the "
        "cerebral circulation [1,2].",
    )
    add_paragraph(
        doc,
        "Transfer function analysis (TFA) expresses that relationship in the "
        "frequency domain. The transfer-function gain describes the relative "
        "size of the output oscillation for a given input oscillation, while "
        "phase describes their timing. Magnitude-squared coherence describes "
        "how consistently the two signals follow a linear frequency-specific "
        "relationship. CARNet recommendations have improved the consistency "
        "of this analysis by defining important choices such as preprocessing, "
        "windowing, frequency bands, and coherence reporting [2].",
    )
    add_paragraph(
        doc,
        "Most applied TFA studies use one input—usually MAP—and one output—"
        "usually middle cerebral artery CBFV. This SISO model is easy to "
        "calculate and interpret. Its simplicity is valuable, but it also "
        "makes an important assumption: other variables that affect CBFV are "
        "either negligible, unrelated to MAP, or absorbed harmlessly into the "
        "residual. That assumption is not always defensible.",
    )

    add_section(doc, "1.2 Why end-tidal carbon dioxide belongs in the question", 2)
    add_paragraph(
        doc,
        "Arterial carbon dioxide is a potent regulator of cerebral vascular "
        "tone [3]. PₑₜCO₂ is an imperfect but practical noninvasive proxy for "
        "arterial carbon dioxide. It changes with breathing depth and timing, "
        "and its spontaneous low-frequency variation can be accompanied by "
        "changes in CBFV. MAP and PₑₜCO₂ can also be statistically related "
        "through respiratory, autonomic, and mechanical pathways. When both "
        "inputs contain shared frequency content, the pairwise MAP–CBFV "
        "relationship may include a carbon-dioxide-associated component.",
    )
    add_paragraph(
        doc,
        "This issue is not new. Panerai and colleagues used a two-input model "
        "of arterial pressure and end-tidal carbon dioxide in healthy adults "
        "and showed that carbon dioxide contributed to CBFV prediction [4]. "
        "Peng and colleagues later demonstrated in the frequency domain that "
        "gas-related inputs can modify the low-frequency pressure transfer "
        "function [5]. Katsogridakis and colleagues showed that adding "
        "physiologically meaningful inputs increased the amount of CBFV "
        "variation represented by a multivariable model and used partial "
        "coherence to distinguish conditional relationships [6]. The 2024 "
        "CARNet time-domain white paper formalized the same two-input logic and "
        "recommended attention to PₑₜCO₂ delay, signal power, sampling, and "
        "regularization [7].",
    )
    add_paragraph(
        doc,
        "Multiple-input nonlinear models have also been used to derive proposed "
        "cerebrovascular “physiomarkers” in normal aging, mild cognitive "
        "impairment, and Alzheimer disease [8–10]. Those studies establish that "
        "multi-input modeling has biological promise. They do not, however, "
        "remove the need to validate a simpler frequency-domain estimator under "
        "the short, noisy recordings that are common in dCA research.",
    )

    add_section(doc, "1.3 The unresolved methodological problem", 2)
    add_paragraph(
        doc,
        "The mathematical MISO solution is compact: at each frequency, solve a "
        "2 × 2 linear system. The difficult part is deciding whether the answer "
        "is reliable. A second input can reduce omitted-input bias when that "
        "input truly affects the output and shares information with the first "
        "input. The same second input can increase variance when its power is "
        "low, when the inputs are nearly collinear, when timing is misaligned, "
        "or when too few independent spectral averages are available. A larger "
        "multiple coherence can also be misleading because adding predictors "
        "will usually improve in-sample fit even when the new pathway estimate "
        "is unstable.",
    )
    add_paragraph(
        doc,
        "The main gap is therefore not the absence of a two-input equation. The "
        "gap is a clear validation and reporting framework for the practical "
        "Welch-based frequency-domain estimator. Researchers need to know how "
        "bias, variance, conditioning, delay, record length, spectral settings, "
        "and regularization interact. They also need evidence that an observed "
        "MISO–SISO difference is linked to the measured PₑₜCO₂–CBFV "
        "relationship rather than to the mere presence of an additional noisy "
        "regressor.",
    )

    add_section(doc, "1.4 Study objective and questions", 2)
    add_paragraph(
        doc,
        "The objective of this approach paper is to develop, validate, and "
        "explain a two-input frequency-domain framework for simultaneous MAP "
        "and PₑₜCO₂ modeling. The paper is organized around four questions:",
    )
    add_numbered(
        doc,
        [
            "Does the unregularized MISO estimator recover known MAP-to-CBFV and "
            "PₑₜCO₂-to-CBFV transfer functions under favorable conditions?",
            "When does adding PₑₜCO₂ reduce error in the MAP pathway compared "
            "with a MAP-only SISO estimate?",
            "Which combinations of low PₑₜCO₂ power, high input coherence, "
            "short duration, timing error, and model misspecification make the "
            "MISO solution unstable or misleading?",
            "Which diagnostics, sensitivity analyses, and reporting rules are "
            "needed before a MISO pathway estimate is interpreted biologically?",
        ],
    )
    add_paragraph(
        doc,
        "We expected MISO to reduce MAP-pathway bias when PₑₜCO₂ had a true "
        "effect on CBFV and shared frequency content with MAP. We expected "
        "little benefit when the PₑₜCO₂ contribution was negligible. We "
        "expected instability when the normalized input spectral matrix was "
        "poorly conditioned, particularly with high input coherence, low "
        "PₑₜCO₂ power, or few Welch averages. Finally, we expected modest "
        "regularization to trade a small amount of bias for a meaningful "
        "reduction in variance under poorly conditioned conditions.",
    )
    add_figure_placeholder(
        doc,
        "1",
        "Conceptual difference between SISO and MISO cerebral hemodynamic models",
        "Panel A: MAP-only SISO (MAP → CBFV). Panel B: PₑₜCO₂-only SISO "
        "(PₑₜCO₂ → CBFV). Panel C: simultaneous MISO with MAP and PₑₜCO₂ "
        "entering one CBFV output and an explicit MAP–PₑₜCO₂ cross-spectral "
        "connection. Panel D: a small visual showing that shared input "
        "information is assigned differently by SISO and MISO.",
        "This figure should allow a reader to understand the entire reason for "
        "the paper before seeing a matrix equation.",
        height_in=1.55,
    )

    # ------------------------------------------------------------------
    # Theory
    # ------------------------------------------------------------------
    add_section(doc, "2. Conceptual and mathematical foundation", 1)
    add_section(doc, "2.1 Signal definitions and scope", 2)
    add_paragraph(
        doc,
        "Let m(t) denote mean-removed MAP, c(t) denote mean-removed PₑₜCO₂, "
        "and y(t) denote mean-removed CBFV. In the empirical analysis, CBFV is "
        "expressed as percent of its recording mean. M(t), C(t), and Y(t) are "
        "their Fourier transforms. The framework assumes that, over the "
        "analyzed record, the measured relationships are approximately linear "
        "and time invariant. This is a working model, not a claim that the "
        "underlying physiology is completely linear, feed-forward, or fixed.",
    )
    add_paragraph(
        doc,
        "We use the cross-spectrum convention S_xy(f) = E[X*(f)Y(f)], where * "
        "denotes complex conjugation. Stating this convention is essential "
        "because reversing the cross-spectrum convention reverses phase sign. "
        "All equations and code in this manuscript follow the same convention.",
    )

    add_section(doc, "2.2 The SISO estimate", 2)
    add_paragraph(
        doc,
        "For a single input x(t) and output y(t), the linear frequency-domain "
        "model is:",
    )
    add_equation(doc, "Y(f) = Hₓᵧ(f)X(f) + E(f)", 1)
    doc.add_page_break()
    add_paragraph(
        doc,
        "If the residual E is uncorrelated with the input, multiplication by "
        "X* and averaging gives the familiar SISO transfer function:",
    )
    add_equation(doc, "Ĥₓᵧˢᴵˢᴼ(f) = Sₓᵧ(f) / Sₓₓ(f)", 2)
    add_paragraph(
        doc,
        "Gain is |Ĥ| and phase is angle(Ĥ). Ordinary magnitude-squared "
        "coherence is:",
    )
    add_equation(doc, "γ²ₓᵧ(f) = |Sₓᵧ(f)|² / [Sₓₓ(f)Sᵧᵧ(f)]", 3)
    add_paragraph(
        doc,
        "These quantities describe a pairwise relationship. They do not by "
        "themselves show that x causes y, and they do not show that an omitted "
        "input is irrelevant.",
    )

    add_section(doc, "2.3 The omitted-input term is visible in the equation", 2)
    add_paragraph(
        doc,
        "Suppose the true linear model has two inputs:",
    )
    add_equation(doc, "Y = HₘM + H꜀C + E", 4)
    add_paragraph(
        doc,
        "A MAP-only SISO analysis still divides S_my by S_mm. Substituting the "
        "two-input model into that calculation gives:",
    )
    add_equation(
        doc,
        "Ĥₘˢᴵˢᴼ = Hₘ + H꜀[Sₘ꜀ / Sₘₘ] + [Sₘₑ / Sₘₘ]",
        5,
    )
    add_paragraph(
        doc,
        "When the residual is unrelated to MAP, the final term is zero in "
        "expectation. The middle term remains whenever PₑₜCO₂ affects CBFV "
        "(H_c ≠ 0) and MAP and PₑₜCO₂ share cross-spectral information "
        "(S_mc ≠ 0). This is the central reason that MISO can matter. It also "
        "shows when MISO should make little difference: if the PₑₜCO₂ pathway "
        "is negligible or the inputs are spectrally unrelated, the omitted-"
        "input term approaches zero.",
    )
    add_callout(
        doc,
        "Core interpretation",
        "A MISO–SISO difference is expected from the mathematics when the "
        "second input both affects the output and shares information with the "
        "first input. The difference is not, by itself, proof that MISO is more "
        "accurate. Known-truth simulation is needed for that claim.",
        kind="info",
    )

    add_section(doc, "2.4 The two-input spectral solution", 2)
    add_paragraph(
        doc,
        "The simultaneous MAP and PₑₜCO₂ model is:",
    )
    add_equation(doc, "Y(f) = Hₘ(f)M(f) + H꜀(f)C(f) + E(f)", 6)
    add_paragraph(
        doc,
        "Multiplying by each conjugated input and taking expectations produces "
        "two spectral normal equations. In matrix form:",
    )
    add_equation(
        doc,
        "[ Sₘₘ  Sₘ꜀ ; S꜀ₘ  S꜀꜀ ][ Hₘ ; H꜀ ] = [ Sₘᵧ ; S꜀ᵧ ]",
        7,
    )
    add_paragraph(
        doc,
        "Using compact notation, S_xxH = S_xy. The estimated transfer-function "
        "vector is obtained by solving:",
    )
    add_equation(doc, "Ĥᴹᴵˢᴼ(f) = Sₓₓ(f) \\ Sₓᵧ(f)", 8)
    # Keep the core implementation excerpt together on a clean page. This
    # explicit break also avoids a LibreOffice rendering quirk in which a
    # naturally paginated table-led page can overlap its header/footer margin.
    doc.add_page_break()
    add_paragraph(
        doc,
        "The backslash symbol in Equation 8 denotes a numerical linear solve, "
        "not division by a scalar. A direct matrix solve is preferred to "
        "forming an explicit inverse because it is clearer and generally more "
        "stable numerically.",
    )
    add_code_block(
        doc,
        "Current two-input solve used in runMISOTFA.m",
        """for k = 1:length(f)
    Sxx = [mapPower(k), mapCo2(k);
           conj(mapCo2(k)), co2Power(k)];
    Sxy = [mapCbv(k); co2Cbv(k)];

    H = Sxx \\ Sxy;
    H_mapcbv(k) = H(1);
    H_co2cbv(k) = H(2);
    conditionNumber(k) = cond(Sxx);
end""",
        "This is an actual shortened extract of the present implementation. "
        "The final paper should link to the exact tagged release.",
    )

    add_section(doc, "2.5 Why input coherence and input power both matter", 2)
    add_paragraph(
        doc,
        "For two inputs, the determinant of the input spectral matrix is:",
    )
    add_equation(doc, "det(Sₓₓ) = SₘₘS꜀꜀ − |Sₘ꜀|²", 9)
    add_paragraph(
        doc,
        "Input coherence is γ²_mc = |S_mc|²/(S_mmS_cc). Substitution gives a "
        "particularly useful identity:",
    )
    add_equation(doc, "det(Sₓₓ) = SₘₘS꜀꜀(1 − γ²ₘ꜀)", 10)
    add_paragraph(
        doc,
        "Equation 10 separates two failure modes. The determinant becomes small "
        "when either input has very low power, or when input coherence approaches "
        "one. In both cases, small spectral-estimation errors can produce large "
        "changes in the pathway coefficients.",
    )
    add_paragraph(
        doc,
        "The raw condition number of S_xx is affected by the different units "
        "and power scales of MAP and PₑₜCO₂. A scale-invariant diagnostic can "
        "be obtained by normalizing the diagonal powers. Define D = "
        "diag(√S_mm, √S_cc) and R = D⁻¹S_xxD⁻¹. R has ones on its diagonal and "
        "the complex input correlation r off the diagonal. Its eigenvalues are "
        "1 − |r| and 1 + |r|. Therefore:",
    )
    add_equation(
        doc,
        "κₙₒᵣₘ(f) = [1 + √γ²ₘ꜀(f)] / [1 − √γ²ₘ꜀(f)]",
        11,
    )
    add_paragraph(
        doc,
        "This normalized condition number measures collinearity but not low "
        "absolute input power. Both κₙₒᵣₘ and the individual power spectra "
        "should be reported. Neither should be converted into an exclusion "
        "threshold until simulation shows what level produces unacceptable "
        "error under the relevant recording conditions.",
    )

    add_section(doc, "2.6 Multiple and partial coherence", 2)
    add_paragraph(
        doc,
        "Multiple coherence summarizes how much output power is represented by "
        "both fitted inputs together:",
    )
    add_equation(
        doc,
        "γ²ᵧ·ₘ꜀(f) = Re{Sₓᵧᴴ(f)Ĥ(f)} / Sᵧᵧ(f)",
        12,
    )
    add_paragraph(
        doc,
        "Because the same record is used to estimate and evaluate the model, "
        "multiple coherence is an in-sample description. It normally rises "
        "when an additional input is included and should not be presented as "
        "out-of-sample validation.",
    )
    add_paragraph(
        doc,
        "Partial coherence asks a different question: how much linear "
        "association remains between one input and the output after the linear "
        "component related to the other input has been removed. For MAP and "
        "CBFV conditional on PₑₜCO₂:",
    )
    add_equation(
        doc,
        "Sₘᵧ·꜀ = Sₘᵧ − (Sₘ꜀S꜀ᵧ / S꜀꜀)",
        13,
    )
    add_equation(
        doc,
        "γ²ₘᵧ|꜀ = |Sₘᵧ·꜀|² / Re{Sₘₘ·꜀Sᵧᵧ·꜀}",
        14,
    )
    add_paragraph(
        doc,
        "Ordinary coherence, partial coherence, and multiple coherence are not "
        "three interchangeable versions of the same metric. Ordinary coherence "
        "is pairwise, partial coherence is conditional, and multiple coherence "
        "describes the combined model. A CARNet ordinary-coherence reference "
        "should not be copied onto partial- or multiple-coherence plots unless "
        "its sampling distribution has been established for those measures.",
    )

    add_section(doc, "2.7 Timing, delay, and phase", 2)
    add_paragraph(
        doc,
        "The measured PₑₜCO₂ signal is not necessarily aligned perfectly with "
        "the cerebral vascular response. Capnograph tubing, breath detection, "
        "circulation time, and physiological response dynamics all contribute "
        "to delay. CARNet notes that a 3–6 second shift is often reasonable "
        "when an experiment-specific delay cannot be estimated reliably [7].",
    )
    add_paragraph(
        doc,
        "A timing error Δτ produces an approximate phase rotation of "
        "2πfΔτ. The same absolute timing error therefore has a larger phase "
        "effect at higher frequency. In a MISO model, misalignment can also "
        "alter how shared MAP–PₑₜCO₂ information is divided between pathways. "
        "Delay should be prespecified or selected using a method that is "
        "evaluated outside the final fit; it should not be chosen only because "
        "it produces the most favorable group result.",
    )
    add_equation(doc, "Δφ(f) = 2πfΔτ", 15)

    add_section(doc, "2.8 Regularization is a bias–variance decision", 2)
    add_paragraph(
        doc,
        "When S_xx is nearly singular, ridge regularization can stabilize the "
        "solution by adding a small positive value to the diagonal. Applying "
        "the same raw value to MAP and PₑₜCO₂ powers is difficult to justify "
        "because their scales differ. A more interpretable approach is to "
        "standardize S_xx first, apply ridge regularization to the normalized "
        "matrix, and then transform the coefficients back to the original "
        "units:",
    )
    add_equation(doc, "β̂λ = (R + λI)⁻¹g ;    Ĥλ = D⁻¹β̂λ", 16)
    add_paragraph(
        doc,
        "Regularization reduces variance but pulls coefficients toward zero. "
        "The preferred λ cannot be chosen from numerical conditioning alone. "
        "It should be evaluated by known-truth error and, where possible, "
        "out-of-sample prediction. The unregularized estimate should remain "
        "available as a reference.",
    )
    add_code_block(
        doc,
        "Standardized ridge MISO sensitivity analysis",
        """% Planned sensitivity implementation
D = diag(sqrt(real(diag(Sxx))));
R = D \\ Sxx / D;
g = D \\ Sxy;

beta = (R + lambda*eye(2)) \\ g;
Hridge = D \\ beta;""",
        "Add this only after testing the transformation with synthetic systems "
        "and verifying the λ grid. Report the unregularized result beside the "
        "regularized sensitivity result.",
        placeholder=True,
    )

    add_section(doc, "2.9 What the model can and cannot mean", 2)
    add_paragraph(
        doc,
        "The MISO coefficients are conditional linear associations in the "
        "measured system. They are not direct measurements of isolated causal "
        "pathways. Feedback can exist between blood pressure, breathing, and "
        "cerebral blood flow. PₑₜCO₂ is not identical to arterial PCO₂. Other "
        "inputs—oxygen, cardiac output, autonomic activity, metabolism, and "
        "neural activity—remain outside the two-input model. The most defensible "
        "language is therefore “MAP-associated” and “PₑₜCO₂-associated” CBFV "
        "dynamics, not “MAP-mediated” and “CO₂-mediated” flow.",
    )
    add_figure_placeholder(
        doc,
        "2",
        "When the two-input solution is identifiable",
        "A two-dimensional map with input coherence on the horizontal axis and "
        "PₑₜCO₂ power or PₑₜCO₂ signal-to-noise ratio on the vertical axis. "
        "Overlay regions of low, intermediate, and high coefficient error from "
        "simulation. Include small marginal panels showing κₙₒᵣₘ and "
        "det(S_xx) behavior.",
        "This should convert the conditioning equations into a practical visual "
        "answer to “when can I trust the MISO estimate?”",
        height_in=1.45,
    )

    # ------------------------------------------------------------------
    # Methods
    # ------------------------------------------------------------------
    add_section(doc, "3. Methods", 1)
    add_section(doc, "3.1 Overall study design", 2)
    add_paragraph(
        doc,
        "The approach is evaluated in three connected stages. Stage 1 is a "
        "known-truth simulation in which the true MAP and PₑₜCO₂ transfer "
        "functions are specified before data are generated. Stage 2 is a set "
        "of stress tests for conditioning, delay, spectral settings, "
        "regularization, and model misspecification. Stage 3 is an empirical "
        "demonstration in NC participants. The simulation establishes accuracy; "
        "the NC analysis establishes practical relevance. The NC data alone are "
        "not treated as proof that one estimator is closer to an unknown "
        "physiological truth.",
    )
    add_callout(
        doc,
        "Analysis-status convention",
        "Methods are written prospectively so the analysis can be implemented "
        "directly. Results that already exist are reported in Section 4. "
        "Unfinished results remain in labeled placeholders. Freeze the final "
        "simulation grid, seed, outcome definitions, and decision rules before "
        "examining the completed simulation results.",
    )

    add_section(doc, "3.2 Known-truth simulation study", 2)
    add_section(doc, "3.2.1 Generating paired input signals", 3)
    add_paragraph(
        doc,
        "Each simulation replicate will generate two stationary, real-valued "
        "input signals with frequency content shaped to resemble resting MAP "
        "and PₑₜCO₂. A shared-noise construction will control the amount of "
        "input coherence. Let u_m, u_c, and u_s be independent zero-mean "
        "Gaussian innovations. After passage through stable shaping filters "
        "F_m and F_c, the inputs will be:",
    )
    add_equation(
        doc,
        "m = F_m{√(1 − r)u_m + √r u_s} ;   c = F_c{√(1 − r)u_c + √r u_s}",
        17,
    )
    add_paragraph(
        doc,
        "With the stated unit-variance construction, the population coherence "
        "is r²; therefore r will be the square root of the target coherence. "
        "Realized band coherence will be "
        "verified after generation; replicates will not be silently discarded "
        "for missing a target. Instead, realized coherence will be retained as "
        "a continuous explanatory variable.",
    )
    add_paragraph(
        doc,
        "A second simulation route will generate Fourier coefficients directly "
        "from a prespecified Hermitian spectral matrix and apply an inverse "
        "Fourier transform. Agreement between the time-domain shared-source "
        "generator and the frequency-domain generator will be used as a code "
        "check, not as two independent biological experiments.",
    )

    add_section(doc, "3.2.2 Known pressure and carbon-dioxide pathways", 3)
    add_paragraph(
        doc,
        "The simulated output will be the sum of a pressure-driven component, a "
        "PₑₜCO₂-driven component, and independent output noise:",
    )
    add_equation(doc, "y(t) = hₘ(t) * m(t) + h꜀(t) * c(t − τ) + e(t)", 18)
    add_paragraph(
        doc,
        "The pressure filter h_m will have a high-pass-like shape over the dCA "
        "range, with stronger damping of slow MAP oscillations than faster "
        "oscillations. The carbon-dioxide filter h_c will emphasize slower "
        "responses. The exact coefficients will be fixed in a simulation "
        "configuration file and displayed in the supplement. A null-carbon-"
        "dioxide condition will set h_c = 0. This condition is important "
        "because it tests whether MISO creates a spurious pathway or needlessly "
        "damages the MAP estimate when the second input has no true effect.",
    )

    add_section(doc, "3.2.3 Simulation factors", 3)
    add_paragraph(
        doc,
        "The main Monte Carlo grid is designed to vary the conditions most "
        "likely to affect the estimator while keeping computation manageable. "
        "Five hundred replicates per main-grid cell are proposed. Monte Carlo "
        "standard errors will be reported for the main performance summaries.",
    )
    add_table(
        doc,
        "Table 1. Proposed known-truth simulation factors.",
        ["Factor", "Main-grid levels", "Question answered"],
        [
            [
                "Input coherence",
                "0, 0.25, 0.50, 0.75, 0.90",
                "How much shared input information can the estimator tolerate?",
            ],
            [
                "PₑₜCO₂ pathway size",
                "None, weak, moderate, strong",
                "When does the second input reduce omitted-input bias?",
            ],
            [
                "Output SNR",
                "0, 5, 10, 20 dB",
                "How does measurement/unmodeled noise affect bias and variance?",
            ],
            [
                "Record duration",
                "5, 10, 20 min",
                "How many data are needed for stable low-frequency estimation?",
            ],
            [
                "PₑₜCO₂ delay",
                "0 s in main grid; 3, 6, 10 s stress test",
                "How sensitive are gain and phase to alignment?",
            ],
            [
                "Estimator",
                "MAP SISO; CO₂ SISO; unregularized MISO; ridge MISO",
                "Which method has the lowest known-truth error?",
            ],
        ],
        [1.35, 2.05, 3.40],
        note=(
            "The final grid should be frozen before results are viewed. If run "
            "time is excessive, reduce Monte Carlo replicates only after using "
            "a precision calculation, not by removing difficult conditions."
        ),
        font_size=8.4,
    )

    add_section(doc, "3.2.4 Secondary stress tests", 3)
    add_paragraph(
        doc,
        "The following stress tests will be run one factor at a time around a "
        "moderate baseline condition, followed by a combined adverse condition:",
    )
    add_bullets(
        doc,
        [
            "Low absolute PₑₜCO₂ power while input coherence is held constant.",
            "A third unmeasured input that affects CBFV and is partly correlated "
            "with MAP or PₑₜCO₂.",
            "Mild nonlinearity, implemented as a quadratic or saturating "
            "PₑₜCO₂ response.",
            "Slow time variation in one pathway coefficient.",
            "Colored rather than white output noise.",
            "PₑₜCO₂ measurement noise added before spectral estimation.",
            "Mismatch between the true PₑₜCO₂ delay and the delay used in the "
            "analysis.",
            "Alternative Welch windows, overlap, detrending, smoothing, and "
            "sampling rates.",
        ],
    )
    add_paragraph(
        doc,
        "These tests are not intended to reproduce every aspect of cerebral "
        "physiology. They test whether the proposed diagnostics fail in "
        "plausible departures from the ideal two-input linear model.",
    )

    add_section(doc, "3.2.5 Simulation outcomes", 3)
    add_paragraph(
        doc,
        "Performance will be evaluated separately for the MAP and PₑₜCO₂ "
        "transfer functions. The primary known-truth outcome will be integrated "
        "squared complex transfer-function error over 0.005–0.20 Hz. This "
        "outcome preserves both magnitude and phase information. Supporting "
        "outcomes will include:",
    )
    add_bullets(
        doc,
        [
            "Frequency- and band-specific gain bias.",
            "Circular phase error.",
            "Root-mean-square error and median absolute error.",
            "Between-replicate coefficient variance.",
            "Error in the null PₑₜCO₂ condition.",
            "Out-of-sample CBFV reconstruction error on an independent "
            "realization generated from the same known system.",
            "Frequency-specific κₙₒᵣₘ, raw condition number, input power, and "
            "the proportion of estimates that are non-finite or extreme.",
            "The relationship between each proposed diagnostic and actual "
            "known-truth error.",
        ],
    )
    add_paragraph(
        doc,
        "The paper will report where MISO improves MAP-pathway error relative "
        "to SISO and where its variance cost is larger than its bias benefit. "
        "A diagnostic will be considered useful only if it predicts actual "
        "estimation error across more than one simulation condition.",
    )
    add_code_block(
        doc,
        "Monte Carlo simulation loop",
        """% Pseudocode to replace with tested MATLAB implementation
for condition = simulationGrid
    for replicate = 1:numReplicates
        [map, co2] = generateCorrelatedInputs(condition);
        cbv = filter(hMap, map) + ...
              filter(hCo2, delayedCo2(co2, condition.delay)) + ...
              generateOutputNoise(condition);

        spectra = estimateWelchSpectra(map, co2, cbv, fs, welch);
        siso = runSISOTFA(spectra, phaseSettings);
        miso = runMISOTFA(spectra, phaseSettings);
        scoreAgainstKnownTruth(siso, miso, hMap, hCo2);
    end
end""",
        "The final excerpt should show the reproducible random seed, "
        "configuration object, and saved condition metadata. The simulation "
        "generator needs unit tests for realized power, coherence, delay, and "
        "known transfer functions.",
        placeholder=True,
    )

    add_section(doc, "3.3 Empirical NC demonstration", 2)
    add_section(doc, "3.3.1 Participants", 3)
    add_paragraph(
        doc,
        "The empirical analysis is a cross-sectional secondary analysis of "
        "resting baseline recordings from 23 normal-control participants. The "
        "approach paper uses only the NC group so that method behavior can be "
        "described without mixing the validation question with a disease-marker "
        "question. All 23 recordings passed the configured requirement of at "
        "least three Welch windows after preprocessing.",
    )
    add_callout(
        doc,
        "Protocol text required from the parent study",
        "Insert the parent study name, recruitment setting, normal-control "
        "definition, inclusion and exclusion criteria, institutional review "
        "board name and approval number, written-consent language, and whether "
        "the 23 participants are independent of any development cohort used to "
        "choose preprocessing or model settings. Do not infer these details "
        "from the spreadsheet.",
    )

    add_section(doc, "3.3.2 Physiological recording", 3)
    add_paragraph(
        doc,
        "Continuous arterial pressure, middle cerebral artery CBFV, PₑₜCO₂, "
        "and electrocardiography were acquired during a resting baseline. The "
        "source recordings were sampled at 500 Hz and later exported as "
        "time-aligned beat-to-beat MAP and CBFV and breath-to-breath PₑₜCO₂. "
        "Cardiac R peaks defined beat intervals.",
    )
    add_callout(
        doc,
        "Hardware and acquisition details required",
        "Add posture, target baseline duration, arterial-pressure device and "
        "calibration, TCD manufacturer and probe frequency, insonated artery "
        "and side, probe fixation, capnograph model and tubing length, room and "
        "breathing conditions, acquisition software, artifact review, and "
        "whether CBFV was unilateral or averaged across sides.",
    )

    add_section(doc, "3.3.3 Preprocessing", 3)
    add_paragraph(
        doc,
        "Rows with non-finite time, MAP, PₑₜCO₂, or CBFV values were removed. "
        "Some PₑₜCO₂ exports contained leading zeros and an initial transition. "
        "The software identified the first nonzero PₑₜCO₂ value, estimated a "
        "large-jump threshold from the upper 10% of consecutive changes, and "
        "removed leading samples until a nonzero value was followed by a change "
        "below that threshold. This rule was used only to remove the start-up "
        "portion and should be inspected visually in every participant.",
    )
    add_paragraph(
        doc,
        "All signals were linearly resampled to 4 Hz. Mean MAP, mean PₑₜCO₂, "
        "and mean CBFV were stored before centering. CBFV was expressed as "
        "percent of the participant’s recording mean:",
    )
    add_equation(doc, "CBFV%(t) = 100 × CBFV(t) / mean[CBFV]", 19)
    add_paragraph(
        doc,
        "The mean was removed from MAP, PₑₜCO₂, and normalized CBFV. "
        "Polynomial detrending was available but disabled in the current "
        "analysis. The final sensitivity analysis will compare no detrending "
        "with linear detrending because low-frequency results can depend on "
        "that decision.",
    )

    add_section(doc, "3.3.4 Welch spectral estimation", 3)
    add_paragraph(
        doc,
        "Auto- and cross-spectra were estimated with Welch’s method [12]. Each "
        "segment used a 128-second Hann window with 50% overlap. At 4 Hz, the "
        "window contained 512 samples, the overlap was 256 samples, and the FFT "
        "length was 512 samples. The frequency spacing was 4/512 = 0.0078125 Hz. "
        "At least three windows were required, corresponding to 256 seconds of "
        "data with the configured overlap.",
    )
    add_paragraph(
        doc,
        "The same three-point kernel [0.25, 0.50, 0.25] was convolved with every "
        "auto- and cross-spectrum. Applying identical linear smoothing to all "
        "spectral terms preserves their alignment before the MISO matrix is "
        "assembled. Results were limited to 0–0.35 Hz, with band analyses "
        "starting at 0.005 Hz.",
    )
    add_code_block(
        doc,
        "Shared spectral estimation and cross-spectrum convention",
        """% MATLAB CPSD output is conjugated to match Sxy = E[X*Y]
[mapCo2Cpsd, f] = cpsd(map, co2, window, overlap, nfft, fs);
[mapCbvCpsd, ~] = cpsd(map, cbv, window, overlap, nfft, fs);

mapCo2 = conj(mapCo2Cpsd);
mapCbv = conj(mapCbvCpsd);

kernel = [0.25 0.50 0.25];
SmapCo2 = conv(mapCo2, kernel, "same");
SmapCbv = conv(mapCbv, kernel, "same");""",
        "This shortened extract documents the phase convention and the rule "
        "that every spectrum receives the same smoothing.",
    )
    add_table(
        doc,
        "Table 2. Current empirical spectral configuration.",
        ["Setting", "Current value", "Planned sensitivity"],
        [
            ["Target sampling rate", "4 Hz", "1 Hz and 4 Hz"],
            ["Window", "128-s Hann", "64, 96, and 128 s"],
            ["Overlap", "50%", "25%, 50%, and 75%"],
            ["Minimum windows", "3", "Report realized count; no silent relaxation"],
            ["FFT length", "Window length", "Unchanged within each window setting"],
            ["Smoothing", "[0.25, 0.50, 0.25]", "None versus current kernel"],
            ["Detrending", "Mean removal only", "None versus linear detrend"],
            ["Analysis range", "0–0.35 Hz", "Primary bands begin at 0.005 Hz"],
        ],
        [1.65, 1.75, 3.40],
        font_size=8.5,
    )

    add_section(doc, "3.3.5 Frequency bands", 3)
    add_paragraph(
        doc,
        "Band values were calculated after each participant’s frequency-domain "
        "model was estimated. Each participant contributed one value per band "
        "to a statistical test; frequency bins were not treated as independent "
        "participants. Arithmetic metrics were averaged across bins. Phase, "
        "when reported, was averaged circularly.",
    )
    add_table(
        doc,
        "Table 3. Prespecified frequency bands.",
        ["Band", "Abbreviation", "Range"],
        [
            ["Very-very-low frequency", "VVLF", "0.005 ≤ f < 0.024 Hz"],
            ["Very-low frequency", "VLF", "0.024 ≤ f < 0.070 Hz"],
            ["Low frequency", "LF", "0.070 ≤ f < 0.200 Hz"],
            ["High frequency", "HF", "0.200 ≤ f ≤ 0.350 Hz"],
        ],
        [2.65, 1.20, 2.95],
        font_size=8.8,
    )

    add_section(doc, "3.3.6 SISO and MISO estimation", 3)
    add_paragraph(
        doc,
        "Separate SISO transfer functions were calculated for MAP-to-CBFV and "
        "PₑₜCO₂-to-CBFV using Equation 2. The same spectra were then used in "
        "the simultaneous MISO solve in Equation 8. Gain was the magnitude of "
        "the complex transfer function. Because CBFV was normalized to percent "
        "of baseline and both inputs were measured in mmHg, gain units were "
        "percentage points of baseline CBFV per mmHg.",
    )
    add_paragraph(
        doc,
        "The present empirical result is unregularized. Raw S_xx condition "
        "numbers were stored at every frequency. The revised analysis will add "
        "κₙₒᵣₘ, input powers, determinant-related diagnostics, and a "
        "prespecified standardized-ridge sensitivity analysis. Regularization "
        "will not replace the main unregularized result unless the known-truth "
        "study shows a clear, reproducible error benefit.",
    )

    add_section(doc, "3.3.7 Empirical surrogate-input test", 3)
    add_paragraph(
        doc,
        "A surrogate test will determine whether the observed MISO–SISO MAP "
        "difference depends on the time-aligned PₑₜCO₂ relationship. For each "
        "participant, PₑₜCO₂ will be circularly shifted by random offsets that "
        "are longer than the model memory and that avoid shifts close to zero "
        "or a full record. Circular shifting preserves the PₑₜCO₂ amplitude "
        "distribution and approximately preserves its power spectrum while "
        "breaking the original alignment with MAP and CBFV.",
    )
    add_paragraph(
        doc,
        "The full spectra and MISO model will be recomputed for each surrogate. "
        "The observed band-specific change in MAP gain will be compared with "
        "the participant-level surrogate distribution. A secondary phase-"
        "randomized surrogate can preserve the PₑₜCO₂ power spectrum more "
        "exactly while destroying its cross-spectral phase relationship. The "
        "two surrogate methods answer related but not identical questions and "
        "will be reported separately if both are used.",
    )
    add_code_block(
        doc,
        "Circular-shift PₑₜCO₂ surrogate",
        """% Planned null analysis
validShiftSamples = makeAllowedShifts(numel(co2), fs, minShiftSeconds);
for b = 1:numSurrogates
    shift = validShiftSamples(randi(numel(validShiftSamples)));
    co2Null = circshift(co2, shift);

    nullSpectra = estimateWelchSpectra(map, co2Null, cbv, fs, welch);
    nullMiso = runMISOTFA(nullSpectra, phaseSettings);
    nullDelta(b,:) = bandGain(nullMiso.map) - bandGain(siso.map);
end""",
        "Use a fixed random seed and save the selected shifts. The surrogate "
        "analysis must recompute all cross-spectra; shifting an already "
        "estimated cross-spectrum is not equivalent.",
        placeholder=True,
    )

    add_section(doc, "3.3.8 Empirical sensitivity analyses", 3)
    add_paragraph(
        doc,
        "The following empirical sensitivity analyses are required before the "
        "methodological conclusions are finalized:",
    )
    add_numbered(
        doc,
        [
            "PₑₜCO₂ alignment: repeat the full analysis with 0-, 3-, and "
            "6-second PₑₜCO₂ shifts, with the primary shift defined before "
            "viewing the comparison.",
            "Welch settings: compare 64-, 96-, and 128-second Hann windows, "
            "retaining the same band definitions and reporting the number of "
            "averages produced by each setting.",
            "Preprocessing: compare mean removal alone with linear detrending, "
            "and compare the current three-bin smoothing with no smoothing.",
            "Sampling: compare 1-Hz and 4-Hz resampling for the 0–0.35-Hz range.",
            "Conditioning: report raw condition number, κₙₒᵣₘ, MAP power, "
            "PₑₜCO₂ power, and determinant-related quantities by frequency.",
            "Regularization: evaluate a prespecified λ grid in the normalized "
            "system and show coefficient paths, known-truth error, and NC "
            "result stability.",
            "Influence: repeat band comparisons after leave-one-participant-out "
            "analysis and report whether any conclusion is driven by one "
            "recording.",
        ],
    )

    add_section(doc, "3.4 Statistical analysis", 2)
    add_paragraph(
        doc,
        "All empirical comparisons are within participant. The primary "
        "empirical family is the MISO-versus-SISO MAP gain difference across "
        "VVLF, VLF, LF, and HF. Paired t tests estimate the mean difference and "
        "95% confidence interval. The standardized paired effect is "
        "d_z = mean(d)/SD(d), where d is MISO minus SISO. Because gain ratios "
        "can be skewed, the median participant-level MISO/SISO ratio and "
        "interquartile range are also reported.",
    )
    add_paragraph(
        doc,
        "Benjamini–Hochberg false-discovery-rate correction is applied across "
        "the four primary MAP bands [13]. The four PₑₜCO₂ gain comparisons "
        "form a separate secondary family. Coherence, phase, conditioning, and "
        "surrogate analyses are supporting or diagnostic unless otherwise "
        "prespecified before rerunning the analysis. Frequency-wise tests are "
        "exploratory and should be corrected across the bins of each curve, "
        "with neighboring-bin structure shown rather than reduced to isolated "
        "significance markers.",
    )
    add_callout(
        doc,
        "Paper-specific multiplicity rerun",
        "The current exported workbook adjusted model-comparison P values across "
        "the combined NC/MCI analysis. The values in this draft’s NC-only tables "
        "use the exported raw P values and a paper-specific four-band BH "
        "recalculation. The final results must be regenerated directly by the "
        "NC-only approach-paper analysis script so the analysis family is "
        "documented in code.",
    )

    add_section(doc, "3.5 Software, testing, and reproducibility", 2)
    add_paragraph(
        doc,
        "The current analysis was implemented in MATLAB R2025b. A shared "
        "preprocessing and spectral-estimation pipeline supplies both SISO and "
        "MISO models, preventing hidden differences in windowing or smoothing. "
        "The repository contains tests for preprocessing, Welch spectra, SISO "
        "and MISO calculations, circular phase, band aggregation, statistics, "
        "workflow settings, and Excel export.",
    )
    add_paragraph(
        doc,
        "The validation release should add unit tests for the simulation "
        "generator, the normalized conditioning identity in Equation 11, "
        "regularized coefficient transformations, delay handling, and "
        "surrogate generation. Every simulation condition should be saved with "
        "its random seed, true transfer functions, realized power and "
        "coherence, estimator settings, and software commit.",
    )
    add_callout(
        doc,
        "Repository information required at submission",
        "Add the public repository URL, archived release/DOI, commit identifier, "
        "MATLAB toolbox requirements, run instructions, synthetic example "
        "dataset, and a machine-readable table of every figure and table value.",
    )

    # ------------------------------------------------------------------
    # Results
    # ------------------------------------------------------------------
    add_section(doc, "4. Results", 1)
    add_section(doc, "4.1 Known-truth estimator validation", 2)
    add_callout(
        doc,
        "RESULTS PENDING — do not convert the planned findings into claims",
        "Insert the simulation results here after the code, grid, and random "
        "seed are frozen. At minimum, report MAP and PₑₜCO₂ complex error, "
        "gain bias, phase error, variance, out-of-sample reconstruction error, "
        "and failure rate across input coherence, pathway size, SNR, and "
        "duration. Explicitly identify conditions where SISO is preferable, "
        "where unregularized MISO is preferable, and where ridge MISO is "
        "preferable.",
    )
    add_paragraph(
        doc,
        "Suggested results structure: first confirm recovery in the favorable "
        "condition; second show omitted-input bias as PₑₜCO₂ contribution and "
        "input coherence rise; third show the variance boundary created by high "
        "coherence and low PₑₜCO₂ power; fourth compare unregularized and ridge "
        "MISO; fifth show whether κₙₒᵣₘ and power predict actual error.",
    )
    add_figure_placeholder(
        doc,
        "3",
        "Known-truth performance of SISO and MISO",
        "Rows for MAP and PₑₜCO₂ pathways. Columns for gain bias, phase error, "
        "and integrated complex error. Use heat maps across input coherence and "
        "PₑₜCO₂ pathway size, with separate panels for duration or SNR. Show "
        "SISO, unregularized MISO, and ridge MISO on the same color scale.",
        "This is the central validation figure and should show both the benefit "
        "and the failure boundary of MISO.",
        height_in=1.65,
    )
    add_figure_placeholder(
        doc,
        "4",
        "Conditioning and regularization performance",
        "Panel A: known-truth error versus κₙₒᵣₘ, colored by PₑₜCO₂ power. "
        "Panel B: error versus raw condition number to show the effect of scale. "
        "Panel C: coefficient error across the λ grid. Panel D: out-of-sample "
        "error across λ. Mark the unregularized solution.",
        "The figure should justify any recommended diagnostic or regularization "
        "range rather than relying on an arbitrary condition-number cutoff.",
        height_in=1.50,
    )

    add_section(doc, "4.2 NC participant and recording characteristics", 2)
    add_paragraph(
        doc,
        "The empirical demonstration included 23 NC participants. Mean age was "
        "66.15 ± 7.14 years, and 13 of 23 participants (56.5%) were female. "
        "Recordings contained modest spontaneous PₑₜCO₂ variation and were "
        "approximately five minutes long. Most recordings produced three "
        "128-second Welch windows.",
    )
    add_table(
        doc,
        "Table 4. NC participant and recording characteristics.",
        ["Characteristic", "NC value"],
        [
            ["Participants", "23"],
            ["Age, years", "66.15 ± 7.14 [55.98–80.53]"],
            ["Female sex", "13/23 (56.5%)"],
            ["Education, years", "16.22 ± 2.37 [12–20]"],
            ["MMSE, points", "29.17 ± 0.94 [27–30]"],
            ["MoCA, points", "27.35 ± 1.85 [24–30]"],
            ["Mean MAP, mmHg", "90.73 ± 14.47 [68.64–116.97]"],
            ["Within-recording MAP SD, mmHg", "4.11 ± 1.14 [2.31–7.19]"],
            ["Mean PₑₜCO₂, mmHg", "36.68 ± 3.46 [29.79–42.50]"],
            ["Within-recording PₑₜCO₂ SD, mmHg", "1.41 ± 0.68 [0.58–3.17]"],
            ["Mean CBFV, cm/s", "46.43 ± 11.75 [31.98–81.99]"],
            ["Usable duration, s", "297.35 ± 13.02 [270.5–347.8]"],
            ["Welch windows", "3.04 ± 0.21 [3–4]"],
        ],
        [3.95, 2.85],
        note=(
            "Values are mean ± SD [minimum–maximum] unless otherwise stated. "
            "MMSE, Mini-Mental State Examination; MoCA, Montreal Cognitive "
            "Assessment."
        ),
        font_size=8.7,
    )

    add_section(doc, "4.3 Model dependence of MAP gain in NC", 2)
    add_paragraph(
        doc,
        "Adding PₑₜCO₂ changed the estimated MAP-to-CBFV gain primarily at the "
        "slowest frequencies. In VVLF, mean MISO MAP gain was 1.092 compared "
        "with 1.332 for SISO. The mean paired difference was −0.240 percentage "
        "points of baseline CBFV per mmHg (95% CI, −0.384 to −0.096; raw "
        "P = 0.00226; four-band BH-adjusted P = 0.00904). The standardized "
        "paired effect was d_z = −0.72. The median participant-level MISO/SISO "
        "ratio was 0.837 (IQR, 0.647–0.980), corresponding to a typical "
        "reduction of about 16%.",
    )
    add_paragraph(
        doc,
        "In LF, MISO MAP gain was slightly higher than SISO gain. The mean "
        "difference was 0.065 (95% CI, 0.010 to 0.120; raw P = 0.0232; "
        "BH-adjusted P = 0.0464; d_z = 0.51). The median ratio was 1.028, so the "
        "typical change was only about 3% despite statistical support. VLF and "
        "HF MAP gain did not show supported model differences. The combination "
        "of effect size, confidence interval, and ratio is important here: the "
        "VVLF result was both statistically and practically larger, whereas "
        "the LF result was small.",
    )
    add_table(
        doc,
        "Table 5. Primary NC comparison of MISO and SISO MAP gain.",
        ["Band", "MISO", "SISO", "Difference (95% CI)", "d_z", "Median ratio", "BH P"],
        [
            ["VVLF", "1.092", "1.332", "−0.240 (−0.384, −0.096)", "−0.72", "0.837", "0.0090"],
            ["VLF", "1.288", "1.270", "0.018 (−0.060, 0.096)", "0.10", "0.997", "0.6370"],
            ["LF", "2.111", "2.046", "0.065 (0.010, 0.120)", "0.51", "1.028", "0.0464"],
            ["HF", "2.343", "2.331", "0.012 (−0.036, 0.060)", "0.11", "1.014", "0.6370"],
        ],
        [0.60, 0.70, 0.70, 2.15, 0.60, 1.05, 0.70],
        note=(
            "Gain units are percentage points of baseline CBFV per mmHg. "
            "Difference is MISO minus SISO. BH adjustment shown here is across "
            "the four NC MAP-gain bands. Final values should be regenerated by "
            "the NC-only analysis script."
        ),
        font_size=8.1,
    )
    add_figure_placeholder(
        doc,
        "5",
        "NC MAP gain: frequency-wise and participant-level MISO–SISO comparison",
        "Panel A: group mean MISO and SISO MAP gain curves with 95% confidence "
        "bands, not only SD. Panel B: paired MISO-minus-SISO difference curve "
        "with a zero line and prespecified band boundaries. Panels C–F: paired "
        "participant plots for VVLF, VLF, LF, and HF, ordered by SISO value. "
        "Annotate mean difference, CI, d_z, median ratio, raw P, and BH P.",
        "This should be the main NC figure. It shows where the model changes the "
        "MAP estimate and whether the group result is shared across participants.",
        height_in=1.65,
    )

    add_section(doc, "4.4 PₑₜCO₂ gain as a supporting model diagnostic", 2)
    add_paragraph(
        doc,
        "The larger model-related changes occurred in the PₑₜCO₂ pathway. "
        "Conditional MISO PₑₜCO₂ gain was lower than PₑₜCO₂-only SISO gain "
        "in LF by −0.737 (95% CI, −1.067 to −0.407; secondary-family BH "
        "P = 0.000259; median ratio, 0.662) and in HF by −2.250 (95% CI, "
        "−3.145 to −1.354; BH P = 0.000127; median ratio, 0.523). VVLF and VLF "
        "differences were not supported after the four-band secondary correction.",
    )
    add_paragraph(
        doc,
        "These results show that a PₑₜCO₂-only SISO coefficient is not "
        "interchangeable with a conditional PₑₜCO₂ coefficient. They do not "
        "establish that the high-frequency MISO PₑₜCO₂ estimate is accurate. "
        "PₑₜCO₂ ordinary coherence was low outside the slowest band, and raw "
        "conditioning became worse at some high-frequency bins. The PₑₜCO₂ "
        "result is therefore most useful as evidence of model dependence and as "
        "a prompt for validation, not as the primary physiological conclusion.",
    )
    add_table(
        doc,
        "Table 6. Secondary NC comparison of MISO and SISO PₑₜCO₂ gain.",
        ["Band", "MISO", "SISO", "Difference (95% CI)", "d_z", "Median ratio", "BH P"],
        [
            ["VVLF", "4.127", "4.370", "−0.243 (−0.675, 0.189)", "−0.24", "0.957", "0.2559"],
            ["VLF", "2.684", "3.170", "−0.486 (−1.087, 0.116)", "−0.35", "0.816", "0.1444"],
            ["LF", "1.761", "2.498", "−0.737 (−1.067, −0.407)", "−0.97", "0.662", "0.00026"],
            ["HF", "2.695", "4.944", "−2.250 (−3.145, −1.354)", "−1.09", "0.523", "0.00013"],
        ],
        [0.60, 0.70, 0.70, 2.15, 0.60, 1.05, 0.70],
        note=(
            "BH adjustment shown here is across the four NC PₑₜCO₂-gain "
            "bands as a secondary family. High-frequency interpretation must "
            "be paired with power, coherence, and conditioning diagnostics."
        ),
        font_size=8.1,
    )

    add_section(doc, "4.5 Coherence structure and in-sample representation", 2)
    add_paragraph(
        doc,
        "MAP–PₑₜCO₂ input coherence was modest on average and decreased from "
        "0.314 in VVLF to approximately 0.21 in LF and HF. Conditional "
        "PₑₜCO₂ partial coherence was larger than conditional MAP partial "
        "coherence in VVLF (0.409 versus 0.335). From VLF through HF, MAP "
        "partial coherence was larger. Multiple coherence ranged from 0.645 to "
        "0.733 across bands. These values describe how the measured signals "
        "are represented in the fitted data; they are not an out-of-sample "
        "accuracy test.",
    )
    add_table(
        doc,
        "Table 7. Descriptive NC coherence values.",
        ["Metric", "VVLF", "VLF", "LF", "HF"],
        [
            ["MAP–PₑₜCO₂ input coherence", "0.314", "0.250", "0.216", "0.209"],
            ["MISO MAP partial coherence", "0.335", "0.560", "0.592", "0.656"],
            ["MISO PₑₜCO₂ partial coherence", "0.409", "0.346", "0.221", "0.225"],
            ["MISO multiple coherence", "0.645", "0.711", "0.681", "0.733"],
            ["SISO MAP ordinary coherence", "0.388", "0.542", "0.584", "0.647"],
            ["SISO PₑₜCO₂ ordinary coherence", "0.456", "0.300", "0.198", "0.209"],
        ],
        [3.20, 0.90, 0.90, 0.90, 0.90],
        note=(
            "Ordinary, partial, and multiple coherence answer different "
            "questions. No ordinary-coherence threshold is applied to partial "
            "or multiple coherence."
        ),
        font_size=8.4,
    )

    add_section(doc, "4.6 Numerical conditioning in the current implementation", 2)
    add_paragraph(
        doc,
        "The raw input spectral matrix was more poorly conditioned at some "
        "high-frequency points. The median raw HF condition number was 10.9. "
        "Values greater than 100 occurred in 60 of 437 HF frequency-by-subject "
        "observations, and the maximum was approximately 7,467. High values "
        "tended to occur when PₑₜCO₂ power was low or when MAP and PₑₜCO₂ "
        "power scales were strongly imbalanced.",
    )
    add_paragraph(
        doc,
        "Because the raw condition number depends on scale, these values do not "
        "prove severe physiological collinearity. They do establish that some "
        "high-frequency coefficients may be sensitive to small spectral "
        "changes. The scale-invariant κₙₒᵣₘ analysis and known-truth error "
        "calibration are pending.",
    )

    add_section(doc, "4.7 Surrogate, delay, and robustness results", 2)
    add_callout(
        doc,
        "RESULTS PENDING",
        "Insert the circular-shift/phase-randomized surrogate results, PₑₜCO₂ "
        "delay analysis, Welch/preprocessing sensitivity, standardized-ridge "
        "coefficient paths, and leave-one-participant-out influence results. "
        "Show estimates and uncertainty, not only whether a P value crosses "
        "0.05.",
    )
    add_figure_placeholder(
        doc,
        "6",
        "NC robustness and falsification analyses",
        "Panel A: observed MAP-gain changes against circular-shift PₑₜCO₂ null "
        "distributions. Panel B: MAP-gain difference at 0-, 3-, and 6-second "
        "PₑₜCO₂ shifts. Panel C: effect estimates across Welch windows, "
        "detrending, smoothing, and sampling rate. Panel D: coefficient paths "
        "across λ. Use VVLF as the primary panel and show all bands in the "
        "supplement.",
        "This figure decides whether the NC model difference is linked to the "
        "measured second input and whether it survives reasonable analysis choices.",
        height_in=1.55,
    )

    # ------------------------------------------------------------------
    # Discussion
    # ------------------------------------------------------------------
    add_section(doc, "5. Discussion", 1)
    add_section(doc, "5.1 Main message", 2)
    add_paragraph(
        doc,
        "The main methodological message is simple: MISO is a conditional "
        "estimator, not an automatic improvement over SISO. A MAP-only transfer "
        "function is a valid description of the pairwise MAP–CBFV relationship. "
        "It is not necessarily an isolated pressure pathway. Adding PₑₜCO₂ "
        "can remove a shared component from the MAP coefficient, but the added "
        "coefficient must be supported by sufficient independent input "
        "information.",
    )
    add_paragraph(
        doc,
        "The current NC results show that this distinction is practically "
        "relevant. MISO MAP gain was lower than SISO gain in VVLF, with a "
        "moderate-to-large paired effect and a median ratio near 0.84. At VLF "
        "and HF, the estimates were similar. The LF difference reached the "
        "paper-specific false-discovery-rate criterion but was small in "
        "relative terms. This pattern argues against a blanket statement that "
        "all MAP gain estimates change once carbon dioxide is included. The "
        "effect is frequency dependent.",
    )
    add_paragraph(
        doc,
        "At the same time, the present NC result is not sufficient to declare "
        "MISO more accurate. The true human transfer functions are unknown, "
        "most records supplied only three Welch windows, no PₑₜCO₂ delay was "
        "applied, and some high-frequency matrices were poorly conditioned. "
        "The final strength of the paper will come from the known-truth and "
        "falsification analyses, not from the significance of one empirical band.",
    )

    add_section(doc, "5.2 How the NC result follows from the model", 2)
    add_paragraph(
        doc,
        "Equation 5 explains the direction of the main NC result. A MAP-only "
        "SISO coefficient contains an additional term proportional to the "
        "PₑₜCO₂ pathway and the MAP–PₑₜCO₂ cross-spectrum. Removing that "
        "shared term can lower the conditional MAP coefficient. The effect "
        "should be strongest where PₑₜCO₂ has meaningful power and coupling "
        "to CBFV, which is most plausible at slow frequencies. The NC data are "
        "consistent with that expectation: PₑₜCO₂ partial coherence was "
        "largest in VVLF, and the clearest MAP model difference occurred in the "
        "same band.",
    )
    add_paragraph(
        doc,
        "The large reductions in PₑₜCO₂ gain at LF and HF should be interpreted "
        "more cautiously. A PₑₜCO₂-only SISO model can assign shared MAP "
        "information to PₑₜCO₂, so a lower conditional coefficient is "
        "mathematically plausible. However, low PₑₜCO₂ power and coherence "
        "also make both SISO and MISO PₑₜCO₂ coefficients less stable. A large "
        "model difference can therefore be caused by successful separation, "
        "numerical sensitivity, or both. The simulation and regularization "
        "analyses are needed to separate those explanations.",
    )

    add_section(doc, "5.3 Relationship to prior work", 2)
    add_paragraph(
        doc,
        "This study builds on, rather than replaces, prior multivariable "
        "cerebral hemodynamic modeling. Panerai and colleagues showed in healthy "
        "participants that including end-tidal carbon dioxide improved a "
        "two-input dynamic model, while the pressure response changed less "
        "strongly than the carbon-dioxide response [4]. Peng and colleagues "
        "showed that carbon dioxide and oxygen reactivity can modify the "
        "low-frequency transfer function previously attributed to pressure "
        "alone [5]. Katsogridakis and colleagues demonstrated that additional "
        "physiologically meaningful inputs increase represented CBFV variation "
        "and that partial coherence can separate conditional associations [6].",
    )
    add_paragraph(
        doc,
        "The CARNet time-domain white paper provides a direct conceptual basis "
        "for the present two-input model. It describes MAP and PₑₜCO₂ as "
        "parallel inputs to CBFV, notes that the second input can reduce bias in "
        "the pressure response, recommends attention to a PₑₜCO₂ delay, and "
        "discusses regularization for ill-conditioned estimation [7]. The "
        "present paper’s novelty is therefore not the existence of two-input "
        "modeling. Its intended contribution is a transparent validation of the "
        "frequency-domain 2 × 2 spectral solve under finite records, together "
        "with practical diagnostics and a reproducible reporting framework.",
    )
    add_paragraph(
        doc,
        "The MCI physiomarker literature is also important context. Marmarelis "
        "and colleagues used subject-specific nonlinear or principal-dynamic-"
        "mode models to estimate pressure autoregulation and carbon-dioxide "
        "vasomotor reactivity [8–10]. More recent PRBS gas challenges provide "
        "stronger broadband excitation than resting spontaneous fluctuations "
        "and may expose dynamics that short baseline records cannot resolve "
        "[11]. Those studies motivate a future marker paper. The current paper "
        "has a narrower goal: establish that the simple MISO estimator is "
        "mathematically and numerically trustworthy before using its "
        "coefficients as candidate biological markers.",
    )

    add_section(doc, "5.4 Why multiple coherence is not enough", 2)
    add_paragraph(
        doc,
        "A two-input model will generally have equal or higher in-sample "
        "multiple coherence than either pairwise SISO coherence. That increase "
        "does not show that both coefficients are accurate. A nearly collinear "
        "second input can improve in-sample representation while making the "
        "individual coefficients unstable. Similarly, a low-noise but "
        "physiologically irrelevant input can absorb residual variation. "
        "Validation must therefore evaluate coefficient error against known "
        "truth and reconstruction error on independent data.",
    )
    add_paragraph(
        doc,
        "The same caution applies to partial coherence. It is a useful "
        "conditional association measure, but its ordinary-coherence reference "
        "threshold is not automatically the same as the CARNet SISO threshold. "
        "Partial coherence should be reported as a distinct quantity with its "
        "own uncertainty and simulation behavior.",
    )

    add_section(doc, "5.5 A practical decision framework", 2)
    add_paragraph(
        doc,
        "A MISO result should be interpreted only after answering five "
        "questions:",
    )
    add_numbered(
        doc,
        [
            "Was there enough MAP and PₑₜCO₂ power in the frequency range of "
            "interest?",
            "Were the inputs sufficiently distinct, as shown by input coherence "
            "and κₙₒᵣₘ?",
            "Was PₑₜCO₂ timing defined and tested?",
            "Was the result stable to reasonable spectral settings, influence "
            "analysis, and prespecified regularization?",
            "Did known-truth or independent-data testing show that the added "
            "input improves the quantity being claimed?",
        ],
    )
    add_paragraph(
        doc,
        "If the answer to the first two questions is no, a pathway coefficient "
        "may be mathematically computable but scientifically weak. If the "
        "answer to the third or fourth question is unknown, the result should "
        "be labeled sensitive or preliminary. If the answer to the fifth "
        "question is no, higher in-sample coherence should not be used as a "
        "substitute.",
    )

    add_table(
        doc,
        "Table 8. Minimum reporting checklist for two-input cerebral TFA.",
        ["Domain", "Minimum item to report"],
        [
            ["Signals", "MAP, PₑₜCO₂, CBFV definitions, units, and measurement devices"],
            ["Timing", "Alignment method, PₑₜCO₂ delay, and tubing/acquisition details"],
            ["Preprocessing", "Artifact handling, resampling, normalization, detrending"],
            ["Spectra", "Window, overlap, FFT length, smoothing, number of averages"],
            ["Bands", "Exact edges and inclusion rules"],
            ["Convention", "Cross-spectrum definition and phase sign"],
            ["Inputs", "MAP and PₑₜCO₂ power plus input coherence"],
            ["Conditioning", "Raw and normalized condition numbers; determinant-related diagnostics"],
            ["Estimator", "SISO equations, MISO solve, and any regularization"],
            ["Coherence", "Ordinary, partial, and multiple coherence kept distinct"],
            ["Uncertainty", "Participant-level estimates, CIs, effects, multiplicity family"],
            ["Validation", "Known-truth error and/or independent-data prediction"],
            ["Sensitivity", "Delay, spectral settings, regularization, and influence analyses"],
            ["Reproducibility", "Code release, commit, configuration, and synthetic example"],
        ],
        [1.35, 5.45],
        font_size=8.4,
    )

    add_section(doc, "5.6 Strengths", 2)
    add_paragraph(
        doc,
        "The proposed study has several strengths. First, it derives the "
        "estimator from the same cross-spectrum convention used in code. "
        "Second, SISO and MISO receive identical preprocessed signals and "
        "spectra, so their differences arise from model structure rather than "
        "hidden preprocessing differences. Third, the known-truth design "
        "separates bias from variance and includes a null second-input "
        "condition. Fourth, conditioning is separated into scale-invariant "
        "collinearity and absolute input power. Fifth, regularization is treated "
        "as a testable bias–variance choice. Sixth, the surrogate-input analysis "
        "provides a direct falsification test in real recordings. Finally, the "
        "empirical analysis reports effect sizes and ratios in addition to "
        "P values.",
    )

    add_section(doc, "5.7 Limitations", 2)
    add_paragraph(
        doc,
        "Several limitations remain. The simulation will simplify a system "
        "that is physiologically nonlinear, time varying, and affected by "
        "feedback. No finite simulation grid can prove validity in every "
        "setting. The empirical recordings are short and mostly provide three "
        "Welch windows, limiting spectral precision and preventing a strong "
        "within-record training/test split. Repeated or longer recordings would "
        "provide a better empirical reliability test [14].",
    )
    add_paragraph(
        doc,
        "PₑₜCO₂ is a surrogate for arterial carbon dioxide and is sensitive to "
        "ventilation, dead space, sampling location, and delay. Transcranial "
        "Doppler measures velocity rather than volumetric flow and assumes that "
        "the insonated artery diameter is sufficiently stable [15]. The model "
        "contains only two inputs and cannot exclude confounding by oxygen, "
        "cardiac output, autonomic activity, or neural-metabolic coupling. The "
        "current empirical analysis has not yet applied a PₑₜCO₂ delay, "
        "normalized conditioning, ridge sensitivity, or surrogate test.",
    )
    add_paragraph(
        doc,
        "Finally, this approach paper uses NC data intentionally. It cannot "
        "establish a disease biomarker or an MCI-specific effect. Those claims "
        "belong in a separate paper after the method, diagnostic criteria, and "
        "candidate marker definitions are fixed.",
    )

    add_section(doc, "5.8 Future work", 2)
    add_paragraph(
        doc,
        "The next methodological step is to complete the known-truth and "
        "robustness analyses and convert their results into prespecified "
        "diagnostic rules. External validation should then use longer or "
        "repeated resting recordings. Controlled pressure oscillations and "
        "broadband PₑₜCO₂ perturbations, including PRBS protocols, would "
        "provide stronger and more separable inputs than spontaneous baseline "
        "variation. These data would allow direct testing of reliability and "
        "out-of-sample prediction.",
    )
    add_paragraph(
        doc,
        "Only after the approach is frozen should the second paper test whether "
        "conditional MAP- or PₑₜCO₂-associated metrics relate to MCI, "
        "continuous cognitive performance, vascular risk, imaging markers, or "
        "disease progression. Separating these papers reduces analytic "
        "flexibility: the marker study will inherit an estimator and quality-"
        "control framework that were established without optimizing a group "
        "difference.",
    )

    # ------------------------------------------------------------------
    # Conclusion and declarations
    # ------------------------------------------------------------------
    add_section(doc, "6. Conclusion", 1)
    add_paragraph(
        doc,
        "A two-input transfer function provides a direct way to estimate MAP- "
        "and PₑₜCO₂-associated CBFV dynamics simultaneously. Its value comes "
        "from separating shared input information, but its main risk comes from "
        "trying to separate inputs that are weak or too similar. The NC "
        "demonstration shows that including PₑₜCO₂ can materially lower the "
        "estimated MAP gain in the slowest band, while leaving most higher-"
        "frequency MAP estimates relatively unchanged. This is a meaningful "
        "methodological observation, not yet a complete validation.",
    )
    add_paragraph(
        doc,
        "The approach paper will be strongest if it makes both sides visible: "
        "where MISO corrects SISO omitted-input bias and where MISO becomes "
        "unstable. Known-truth simulation, normalized conditioning, delay "
        "sensitivity, surrogate-input testing, and prespecified regularization "
        "are therefore central parts of the paper rather than optional "
        "supplements. Once those analyses are complete, the framework can "
        "support a separate and more defensible search for biological markers.",
    )

    add_section(doc, "Declarations", 1)
    add_section(doc, "Ethics approval and consent to participate", 2)
    add_paragraph(
        doc,
        "[Insert the PI-approved institutional review board statement, approval "
        "number, and written informed-consent language.]",
        italic=True,
    )
    add_section(doc, "Data availability", 2)
    add_paragraph(
        doc,
        "[State whether de-identified physiological data can be shared, whether "
        "a data-use agreement is required, and where synthetic validation data "
        "will be deposited.]",
        italic=True,
    )
    add_section(doc, "Code availability", 2)
    add_paragraph(
        doc,
        "The MATLAB analysis and simulation code will be made available at "
        "[repository URL and archived release to be added].",
    )
    add_section(doc, "Funding", 2)
    add_paragraph(doc, "[Funding and institutional support to be added.]", italic=True)
    add_section(doc, "Conflicts of interest", 2)
    add_paragraph(doc, "[Confirm with every author.]", italic=True)
    add_section(doc, "Author contributions", 2)
    add_paragraph(
        doc,
        "[Add CRediT roles after the final author list is confirmed.]",
        italic=True,
    )
    add_section(doc, "Acknowledgments", 2)
    add_paragraph(doc, "[Acknowledgments to be added.]", italic=True)

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------
    add_section(doc, "References", 1)
    references = [
        "1. Claassen JAHR, Thijssen DHJ, Panerai RB, Faraci FM. Regulation of "
        "cerebral blood flow in humans: physiology and clinical implications "
        "of autoregulation. Physiological Reviews. 2021;101(4):1487–1559. "
        "doi:10.1152/physrev.00022.2020.",
        "2. Panerai RB, Brassard P, Burma JS, et al.; Cerebrovascular Research "
        "Network. Transfer function analysis of dynamic cerebral "
        "autoregulation: a CARNet white paper 2022 update. Journal of Cerebral "
        "Blood Flow & Metabolism. 2023;43(1):3–25. "
        "doi:10.1177/0271678X221119760.",
        "3. Hoiland RL, Fisher JA, Ainslie PN. Regulation of the cerebral "
        "circulation by arterial carbon dioxide. Comprehensive Physiology. "
        "2019;9(3):1101–1154. doi:10.1002/cphy.c180021.",
        "4. Panerai RB, Simpson DM, Deverson ST, Mahony P, Hayes P, Evans DH. "
        "Multivariate dynamic analysis of cerebral blood flow regulation in "
        "humans. IEEE Transactions on Biomedical Engineering. "
        "2000;47(3):419–423. doi:10.1109/10.827312.",
        "5. Peng T, Rowley AB, Ainslie PN, Poulin MJ, Payne SJ. Multivariate "
        "system identification for cerebral autoregulation. Annals of "
        "Biomedical Engineering. 2008;36(2):308–320. "
        "doi:10.1007/s10439-007-9412-9.",
        "6. Katsogridakis E, Simpson DM, Bush G, et al. Revisiting the frequency "
        "domain: the multiple and partial coherence of cerebral blood flow "
        "velocity in the assessment of dynamic cerebral autoregulation. "
        "Physiological Measurement. 2016;37(7):1056–1073. "
        "doi:10.1088/0967-3334/37/7/1056.",
        "7. Kostoglou K, Bello-Robles F, Brassard P, et al. Time-domain methods "
        "for quantifying dynamic cerebral blood flow autoregulation: review and "
        "recommendations. A white paper from the Cerebrovascular Research "
        "Network (CARNet). Journal of Cerebral Blood Flow & Metabolism. "
        "2024;44(9):1480–1514. doi:10.1177/0271678X241249276.",
        "8. Marmarelis VZ, Shin DC, Orme ME, Zhang R. Model-based physiomarkers "
        "of cerebral hemodynamics in patients with mild cognitive impairment. "
        "Medical Engineering & Physics. 2014;36(5):628–637. "
        "doi:10.1016/j.medengphy.2014.02.025.",
        "9. Marmarelis VZ, Shin DC, Tarumi T, Zhang R. Comparison of model-based "
        "indices of cerebral autoregulation and vasomotor reactivity using "
        "transcranial Doppler versus near-infrared spectroscopy in patients "
        "with amnestic mild cognitive impairment. Journal of Alzheimer’s "
        "Disease. 2017;56:89–105. doi:10.3233/JAD-161004.",
        "10. Marmarelis VZ, Mitsis GD, Shin DC, Zhang R. Multiple-input nonlinear "
        "modelling of cerebral haemodynamics using spontaneous arterial blood "
        "pressure, end-tidal CO₂ and heart rate measurements. Philosophical "
        "Transactions of the Royal Society A. 2016;374(2067):20150180. "
        "doi:10.1098/rsta.2015.0180.",
        "11. Hashem S, Yamashiro S, Joe E, Chui H, Marmarelis V. PRBS gas "
        "challenges reveal impaired chemoreflex and cholinergic dynamics in "
        "MCI. Annals of Biomedical Engineering. 2026;54:2497–2507. "
        "doi:10.1007/s10439-026-04213-7.",
        "12. Welch PD. The use of fast Fourier transform for the estimation of "
        "power spectra: a method based on time averaging over short, modified "
        "periodograms. IEEE Transactions on Audio and Electroacoustics. "
        "1967;15(2):70–73. doi:10.1109/TAU.1967.1161901.",
        "13. Benjamini Y, Hochberg Y. Controlling the false discovery rate: a "
        "practical and powerful approach to multiple testing. Journal of the "
        "Royal Statistical Society Series B. 1995;57(1):289–300. "
        "doi:10.1111/j.2517-6161.1995.tb02031.x.",
        "14. Elting JW, Sanders ML, Panerai RB, et al. Assessment of dynamic "
        "cerebral autoregulation in humans: is reproducibility dependent on "
        "blood pressure variability? PLOS ONE. 2020;15(1):e0227651. "
        "doi:10.1371/journal.pone.0227651.",
        "15. Aaslid R, Markwalder TM, Nornes H. Noninvasive transcranial Doppler "
        "ultrasound recording of flow velocity in basal cerebral arteries. "
        "Journal of Neurosurgery. 1982;57(6):769–774. "
        "doi:10.3171/jns.1982.57.6.0769.",
    ]
    for index, ref in enumerate(references):
        if index == 3:
            # Start the longer reference continuation on a deliberate page so
            # the rendered page retains its header and footer margins.
            doc.add_page_break()
        add_paragraph(doc, ref, style="Reference")

    # ------------------------------------------------------------------
    # Appendices
    # ------------------------------------------------------------------
    doc.add_page_break()
    add_section(doc, "Appendix A. Analysis completion checklist", 1)
    add_table(
        doc,
        "Table A1. Work required before the draft can be treated as submission-ready.",
        ["Item", "Required output", "Status"],
        [
            ["Freeze simulation configuration", "Versioned file with factors, seeds, and true filters", "Pending"],
            ["Run known-truth main grid", "Error tables and Figure 3", "Pending"],
            ["Run stress tests", "Delay, low power, third input, nonlinearity, time variation", "Pending"],
            ["Implement κₙₒᵣₘ", "Frequency-wise and band summaries plus unit test", "Pending"],
            ["Implement ridge sensitivity", "Known-truth and NC coefficient paths", "Pending"],
            ["Run surrogate PₑₜCO₂ null", "Observed-versus-null effect distributions", "Pending"],
            ["Rerun NC-only statistics", "Four-band primary and secondary families in code", "Pending"],
            ["Create main NC figure", "Frequency curves plus paired participant panels", "Pending"],
            ["Confirm protocol details", "PI-approved Methods and ethics text", "Pending"],
            ["Archive code release", "Repository URL, tag, DOI, run instructions", "Pending"],
        ],
        [2.20, 3.75, 0.85],
        font_size=8.3,
    )

    add_section(doc, "Appendix B. Recommended figure and supplement map", 1)
    add_table(
        doc,
        "Table B1. Proposed display items.",
        ["Item", "Role", "Recommended location"],
        [
            ["Figure 1", "Conceptual SISO versus MISO schematic", "Main text"],
            ["Figure 2", "Identifiability map from coherence and power", "Main text"],
            ["Figure 3", "Known-truth estimator performance", "Main text; central figure"],
            ["Figure 4", "Conditioning and regularization", "Main text"],
            ["Figure 5", "NC MAP MISO–SISO comparison", "Main text"],
            ["Figure 6", "NC surrogate and robustness analyses", "Main text or supplement"],
            ["Figure S1", "True simulation filters and input spectra", "Supplement"],
            ["Figure S2", "Full simulation condition grid", "Supplement"],
            ["Figure S3", "PₑₜCO₂ pathway empirical curves", "Supplement"],
            ["Figure S4", "All coherence metrics with correct references", "Supplement"],
            ["Figure S5", "Raw and normalized conditioning by frequency", "Supplement"],
            ["Table S1", "Complete simulation configuration", "Supplement"],
            ["Table S2", "All NC frequency-band estimates", "Supplement"],
            ["Table S3", "Sensitivity-analysis effect estimates", "Supplement"],
        ],
        [1.10, 3.80, 1.90],
        font_size=8.3,
    )

    add_section(doc, "Appendix C. Author decisions that should not be guessed", 1)
    add_bullets(
        doc,
        [
            "Final title, author order, affiliations, and corresponding author.",
            "Target journal and its word, figure, abstract, and reference limits.",
            "Parent protocol description, NC eligibility, ethics number, and "
            "informed-consent wording.",
            "Final PₑₜCO₂ delay rule and whether the unshifted analysis remains "
            "the primary empirical result.",
            "Final simulation grid, Monte Carlo replicates, and primary error "
            "metric.",
            "Whether ridge MISO remains a sensitivity analysis or becomes the "
            "primary estimator after known-truth validation.",
            "Data-sharing limits and the public code/archive location.",
        ],
    )
    add_callout(
        doc,
        "Final drafting principle",
        "Do not fill a missing protocol detail, analysis result, or numerical "
        "threshold because it sounds reasonable. Keep the visible marker until "
        "the source is available or the analysis is complete.",
        kind="info",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
