from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(
    "manuscript/Approach Paper for MISO vs SISO (Changes to Section 1).docx"
)


REPLACEMENTS = {
    "Most applied TFA studies generally use one input, MAP, and one output, middle cerebral artery CBFV. "
    "This SISO model is easy to calculate and interpret. Its simplicity is valuable, but it also makes an "
    "important assumption: other variables that affect CBFV are either negligible, unrelated to MAP, or "
    "absorbed harmlessly into the residual (what is the residual). That assumption is not always defensible.":
    "Most applied TFA studies use one input—usually MAP—and one output—usually middle cerebral artery CBFV. "
    "This SISO model is easy to calculate and interpret. Its simplicity is valuable, but it also makes an "
    "important assumption: other variables that affect CBFV are either negligible, unrelated to MAP, or left "
    "in the residual—the part of the measured CBFV signal that the MAP-only model does not explain. The "
    "residual can contain measurement noise, unmeasured physiological influences, nonlinear behavior, and "
    "ordinary random variation. If one of those influences is related to both MAP and CBFV, leaving it in the "
    "residual can distort the estimated MAP–CBFV relationship.",

    "Arterial carbon dioxide is a potent regulator of cerebral vascular tone [3]. PₑₜCO₂ is an imperfect but "
    "practical noninvasive proxy for arterial carbon dioxide. It changes with breathing depth and timing, and "
    "its spontaneous low-frequency variation can be accompanied by changes in CBFV. MAP and PₑₜCO₂ can also "
    "be statistically related through respiratory, autonomic, and mechanical pathways. When both inputs "
    "contain shared frequency content, the pairwise MAP–CBFV relationship may include a carbon-dioxide-associated "
    "component.":
    "Arterial carbon dioxide is a potent regulator of cerebral vascular tone [3]. PₑₜCO₂ is an imperfect but "
    "practical noninvasive proxy for arterial carbon dioxide. It changes with breathing depth and timing, and "
    "its spontaneous low-frequency variation can be accompanied by changes in CBFV. MAP and PₑₜCO₂ can also "
    "be statistically related through respiratory, autonomic, and mechanical pathways. When MAP and PₑₜCO₂ "
    "tend to rise and fall together at the same frequencies, a pressure-only analysis may attribute part of "
    "the CO₂-associated CBFV response to MAP.",

    "This issue is not new. Panerai and colleagues used a two-input model of arterial pressure and end-tidal "
    "carbon dioxide in healthy adults and showed that carbon dioxide contributed to CBFV prediction [4]. Peng "
    "and colleagues later demonstrated in the frequency domain that gas-related inputs can modify the "
    "low-frequency pressure transfer function (from coherence, remind me or say how exactly here itself) [5]. "
    "Katsogridakis and colleagues showed that adding physiologically meaningful inputs increased the amount of "
    "CBFV variation represented by a multivariable model and used partial coherence to distinguish conditional "
    "relationships (does increasing variance mean even less interpretable? So is that bad? You need to tell the "
    "reader what to think) [6]. The 2024 CARNet time-domain white paper formalized the same two-input logic and "
    "recommended attention to PₑₜCO₂ delay, signal power, sampling, and regularization [7].":
    "This issue is not new. Panerai and colleagues used a two-input model of arterial pressure and end-tidal "
    "carbon dioxide in healthy adults and showed that carbon dioxide contributed to CBFV prediction [4]. Peng "
    "and colleagues examined spontaneous MAP, PₑₜCO₂, and PₑₜO₂ fluctuations. They found that adding the "
    "gas-related inputs significantly increased multiple coherence below 0.05 Hz compared with MAP-only "
    "coherence. They also showed that gas reactivity could alter the low-frequency MAP–CBFV transfer function, "
    "meaning a MAP-only estimate may not represent pressure autoregulation alone [5]. Katsogridakis and "
    "colleagues later found that adding physiologically meaningful inputs allowed a multivariable model to "
    "account for a greater proportion of observed CBFV variation. In this context, greater explained variation "
    "indicated improved model representation; it did not mean that the pathway estimates themselves became "
    "more variable. Their partial-coherence analysis then helped distinguish the relationship of each input "
    "with CBFV after accounting for the other inputs [6]. Together, these findings support multivariable "
    "modeling while also showing why the stability and interpretation of each estimated pathway should be "
    "examined separately. The 2024 CARNet time-domain white paper formalized the same two-input logic and "
    "recommended attention to PₑₜCO₂ delay, signal power, sampling, and regularization [7].",

    "The mathematical MISO solution is compact: at each frequency, solve a 2 × 2 linear system. The difficult "
    "part is deciding whether the answer is reliable. A second input can reduce omitted-input bias when that "
    "input truly affects the output and shares information with the first input. The same second input can "
    "increase variance when its power is low, when the inputs are nearly collinear, when timing is misaligned, "
    "or when too few independent spectral averages are available (this may need to be broken down further. Keep "
    "in mind that the target audience is not engineers but physiologists looking to understanding an engineering "
    "approach to dCA). A larger multiple coherence can also be misleading because adding predictors will usually "
    "improve in-sample fit even when the new pathway estimate is unstable.":
    "The mathematical MISO solution is compact: at each frequency, solve a 2 × 2 linear system. The difficult "
    "part is deciding when the answer is reliable. A second input can reduce omitted-input bias when it truly "
    "affects CBFV and is statistically related to the first input. However, it can also make the estimate less "
    "stable under several circumstances. If PₑₜCO₂ changes very little, the model has little information from "
    "which to estimate the CO₂ pathway. If MAP and PₑₜCO₂ move almost identically, the model has difficulty "
    "determining which input produced the observed CBFV change. A timing error can shift or distort the estimated "
    "gain and phase. Finally, if the recording provides too few data segments for spectral averaging, random "
    "variation in the spectral estimates can have a larger influence on the solution. These problems are "
    "described technically as low input power, input collinearity, timing misalignment, and too few independent "
    "Welch averages. A larger multiple coherence also does not by itself prove that a new pathway estimate is "
    "stable, because a model will generally fit the same data at least as well after another predictor is added.",

    "The main gap is therefore not the absence of a two-input equation. The gap is a clear validation and "
    "reporting framework for the practical Welch-based frequency-domain estimator. Researchers need to know how "
    "bias, variance, conditioning, delay, record length, spectral settings, and regularization interact (this "
    "wording might be a little strong and suggestive that they don’t know but this is also a critical statement "
    "so we need to convey this in a less harsh way). They also need evidence that an observed MISO–SISO difference "
    "is linked to the measured PₑₜCO₂–CBFV relationship rather than to the mere presence of an additional noisy "
    "regressor. (same with the tonality of this sentence but the content is really good)":
    "The main gap is therefore not the absence of a two-input equation. Although these practical factors are "
    "recognized individually, their combined influence on a Welch-based two-input frequency-domain estimator "
    "has not been characterized systematically. A practical validation framework should therefore show how "
    "estimator bias and variability, matrix conditioning, timing, record length, spectral settings, and "
    "regularization affect the resulting pathway estimates. It should also test whether an observed MISO–SISO "
    "difference follows the measured PₑₜCO₂–CBFV relationship and is reduced when that temporal relationship is "
    "deliberately disrupted.",

    "The objective of this approach (I don’t think you need to say approach) paper is to develop, validate, and "
    "explain a two-input frequency-domain framework for simultaneous MAP and PₑₜCO₂ modeling. The paper is "
    "organized around four questions:":
    "The objective of this paper is to develop, validate, and explain a two-input frequency-domain framework "
    "for simultaneous MAP and PₑₜCO₂ modeling. The paper is organized around four questions:",

    "We expected MISO to reduce MAP-pathway bias when PₑₜCO₂ had a true effect on CBFV and shared frequency "
    "content with MAP. We expected little benefit when the PₑₜCO₂ contribution was negligible. We expected "
    "instability when the normalized input spectral matrix was poorly conditioned, particularly with high input "
    "coherence, low PₑₜCO₂ power, or few Welch averages. Finally, we expected modest regularization to trade a "
    "small amount of bias for a meaningful reduction in variance under poorly conditioned conditions "
    "(conditioned conditions sounds a bit weird is there a better word if not it’s fine).":
    "We expected MISO to reduce MAP-pathway bias when PₑₜCO₂ both affected CBFV and fluctuated systematically "
    "with MAP at the same frequencies. We expected little benefit when the PₑₜCO₂ contribution was negligible. "
    "We expected instability when the normalized input spectral matrix was poorly conditioned, particularly with "
    "high input coherence, low PₑₜCO₂ power, or few Welch averages. Finally, we expected modest regularization "
    "to trade a small amount of bias for a meaningful reduction in variance under poorly conditioned "
    "circumstances.",
}


def replace_paragraph_text(paragraph, new_text):
    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)

    paragraph._p.clear_content()
    new_run = paragraph.add_run(new_text)
    if first_rpr is not None:
        new_run._r.insert(0, first_rpr)


document = Document(SOURCE)
remaining = dict(REPLACEMENTS)

for paragraph in document.paragraphs:
    replacement = remaining.pop(paragraph.text, None)
    if replacement is not None:
        replace_paragraph_text(paragraph, replacement)

if remaining:
    missing = "\n\n".join(remaining)
    raise RuntimeError(f"Expected source paragraphs were not found:\n\n{missing}")

document.save(SOURCE)
print(f"Updated {len(REPLACEMENTS)} Section 1 paragraphs in {SOURCE}")
