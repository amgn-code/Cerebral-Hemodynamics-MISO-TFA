from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(
    "manuscript/Approach Paper for MISO vs SISO (Changes to Section 1).docx"
)


REPLACEMENTS = {
    "Most applied TFA studies usually use one input, MAP, and one output middle cerebral artery CBFV. "
    "This SISO model is easy to calculate and interpret. Its simplicity is valuable, but it also makes "
    "an important assumption: other variables that affect CBFV are either negligible, unrelated to MAP, "
    "or left in the residual, the part of the measured CBFV signal that the MAP-only model does not explain. "
    "The residual can contain measurement noise, unmeasured physiological influences, nonlinear behavior, "
    "and ordinary random variation. If one of those influences is related to both MAP and CBFV, leaving it "
    "in the residual can distort the estimated MAP–CBFV relationship. (This now goes too much into residual "
    "here itself. I think it will be more appropriate to get into residual when we start actually talking "
    "about it later in the paper)":
    "Most applied TFA studies usually use one input, MAP, and one output middle cerebral artery CBFV. "
    "This SISO model is easy to calculate and interpret. Its simplicity is valuable, but it assumes that "
    "other variables affecting CBFV are either negligible, unrelated to MAP, or can be left unmodeled. "
    "That assumption is not always defensible.",

    "Multiple-input nonlinear models (such as…) have also been used to derive proposed cerebrovascular "
    "“physiomarkers” in normal aging, mild cognitive impairment, and Alzheimer disease [8–10]. Those "
    "studies establish that multi-input modeling has biological promise. They do not, however, remove the "
    "need to validate a simpler frequency-domain estimator under the short, noisy recordings that are "
    "common in dCA research.":
    "Multiple-input nonlinear models based on principal dynamic modes have also been used to estimate "
    "pressure-related and CO₂-related cerebrovascular responses and derive proposed cerebrovascular "
    "“physiomarkers” in normal aging, mild cognitive impairment, and Alzheimer disease [8–10]. Those "
    "studies establish that multi-input modeling has biological promise. They do not, however, remove the "
    "need to validate a simpler frequency-domain estimator under the short, noisy recordings that are "
    "common in dCA research.",
}


def replace_paragraph_text(paragraph, new_text):
    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)

    paragraph._p.clear_content()
    new_run = paragraph.add_run(new_text)
    if first_rpr is not None:
        new_run._r.insert(0, first_rpr)


document = Document(SOURCE)
remaining = dict(REPLACEMENTS)

for paragraph in document.paragraphs:
    replacement = remaining.pop(paragraph.text, None)
    if replacement is not None:
        replace_paragraph_text(paragraph, replacement)

if remaining:
    missing = "\n\n".join(remaining)
    raise RuntimeError(f"Expected source paragraphs were not found:\n\n{missing}")

document.save(SOURCE)
print(f"Updated {len(REPLACEMENTS)} paragraphs in {SOURCE}")
