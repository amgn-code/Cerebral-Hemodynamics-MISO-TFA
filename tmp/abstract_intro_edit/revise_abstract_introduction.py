from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt


DOCUMENT = Path(__file__).with_name("Abstract & Introduction.docx")


TITLE = (
    "Development and Known-Truth Validation of a Two-Input Transfer "
    "Function Analysis Pipeline for Cerebral Hemodynamics"
)


ABSTRACT = (
    "Transfer Function Analysis (TFA) of dynamic cerebral autoregulation (dCA) "
    "generally treats mean arterial pressure (MAP) as a single input affecting "
    "the output of cerebral blood flow velocity (CBFV). As such, the system "
    "describing dCA is modeled as a single-input-single-output (SISO) system "
    "that provides the complex relationship, including gain and phase, between "
    "MAP and CBFV. However, previous research has shown that other inputs, such "
    "as end-tidal carbon dioxide (PETCO₂), can also influence CBFV. Therefore, "
    "it is important to evaluate a multiple-input-single-output (MISO) model "
    "that accounts for the relationships between both MAP and PETCO₂ and the "
    "output of CBFV. Furthermore, a MISO model can use the shared frequency "
    "content between MAP and PETCO₂ when estimating each input-to-output "
    "relationship, as opposed to two separate SISO models that relate either "
    "MAP or PETCO₂ to CBFV without accounting for the other input. While "
    "previous studies have shown higher multiple coherence with multiple-input "
    "models, this study seeks to determine whether that increased coherence "
    "corresponds to improved pathway accuracy or instead reflects the greater "
    "flexibility of a model with an additional input. In this study, we develop "
    "and validate a transparent MISO TFA pipeline that jointly estimates the "
    "MAP-to-CBFV and PETCO₂-to-CBFV transfer functions. The pipeline provides a "
    "consistent procedure for spectral estimation, the MISO solution, gain and "
    "phase calculation, model diagnostics, statistical comparison, and "
    "reporting. Short recordings, weak PETCO₂ fluctuations, correlated inputs, "
    "measurement noise, timing errors, and poorly conditioned spectral matrices "
    "may offset the conceptual benefit of separating the two pathways. The "
    "pipeline is therefore evaluated using a known-truth simulation framework "
    "to identify the conditions under which two-input MISO transfer function "
    "analysis recovers the true MAP and PETCO₂ pathways more accurately than "
    "separate SISO models. Each simulated family represents a fixed "
    "physiological system, while recording duration, noise, timing alignment, "
    "and estimator settings are varied as observation conditions without "
    "changing the underlying truth. Accuracy will be evaluated using "
    "frequency-resolved gain error, wrapped phase error, normalized complex "
    "error, and a signed model-advantage measure that indicates whether MISO or "
    "SISO has lower error. Resting recordings from cognitively normal adults "
    "will then be placed within the simulated operating space to determine "
    "whether their duration, input relationships, excitation, and conditioning "
    "support interpretable MISO estimates. The purpose of this paper is to "
    "develop a reproducible MISO TFA pipeline, explain the reasoning behind its "
    "major analytical choices, establish the conditions under which the "
    "additional input is justified, determine when SISO remains sufficient, "
    "and identify when neither separated pathway should be interpreted without "
    "additional data."
)


INTRODUCTION_PARAGRAPHS = [
    (
        "The brain requires continuous blood flow even as arterial pressure "
        "changes from moment to moment. Cerebral autoregulation describes the "
        "physiological processes that limit how strongly these pressure changes "
        "alter cerebral perfusion [1]. Dynamic cerebral autoregulation focuses "
        "on responses that occur over seconds to minutes. TFA is commonly used "
        "to study these responses by treating MAP as the input and CBFV as the "
        "output [2]. The resulting transfer function is complex-valued. Gain "
        "describes how strongly changes in MAP appear in CBFV, while phase "
        "describes their timing relationship at each frequency. Coherence is "
        "also used to describe how consistently the two signals are linearly "
        "related."
    ),
    (
        "A MAP-only model is useful, but MAP is not the only input that can "
        "influence CBFV. Carbon dioxide has a strong effect on cerebrovascular "
        "tone, and PETCO₂ provides a practical continuous estimate of carbon "
        "dioxide during resting recordings [3]. When MAP and PETCO₂ share "
        "frequency content, a MAP-only SISO estimate can include part of the "
        "PETCO₂-related response. The same issue applies to a PETCO₂-only SISO "
        "estimate. Therefore, two separate SISO models describe pairwise "
        "relationships, but they do not separate the contribution of one "
        "measured input from the other."
    ),
    (
        "Previous studies have used multivariable models to evaluate MAP and "
        "carbon dioxide together [4-9]. These studies provide a strong basis for "
        "using a two-input model. However, the use of MISO TFA is not yet "
        "organized into one clear and reproducible workflow for resting "
        "cerebral hemodynamic recordings. Studies may use different "
        "preprocessing choices, spectral settings, model solutions, diagnostic "
        "criteria, and reporting methods. In addition, higher multiple "
        "coherence shows that the model fits the measured output more closely, "
        "but it does not show that the separated MAP and PETCO₂ pathways are "
        "more accurate."
    ),
    (
        "The first goal of this paper is therefore to develop and document a "
        "complete MISO TFA pipeline. The pipeline begins with MAP, PETCO₂, and "
        "CBFV signals and applies the same preprocessing and spectral estimation "
        "procedure to each recording. The MAP and PETCO₂ auto-spectra, their "
        "cross-spectrum, and their cross-spectra with CBFV are then used in one "
        "joint MISO solution. The pipeline returns MAP and PETCO₂ gain and "
        "phase, coherence measures, input power, and conditioning diagnostics. "
        "It also provides the statistical comparisons, figures, tables, and "
        "source data needed to understand and reproduce the analysis. Each major "
        "choice will be explained so that the workflow can be inspected and "
        "modified rather than treated as a hidden software procedure."
    ),
    (
        "Adding a second input does not automatically produce a more accurate "
        "model. MISO must estimate two pathway coefficients from the same finite "
        "recording. If MAP and PETCO₂ become too similar within a frequency "
        "region, the input spectral matrix can become poorly conditioned and "
        "small spectral errors can produce large changes in the separated "
        "transfer functions. Short recordings, weak PETCO₂ fluctuations, "
        "measurement noise, signal misalignment, and estimator settings may "
        "also affect this balance. A useful pipeline must therefore report the "
        "conditions that support the MISO estimates, not only the estimates "
        "themselves."
    ),
    (
        "Human recordings cannot show the true MAP-to-CBFV or PETCO₂-to-CBFV "
        "transfer functions. If MISO and SISO produce different results in an "
        "observed recording, that difference alone cannot show which model is "
        "closer to the underlying physiology. A known-truth simulation makes "
        "this comparison possible. MAP and PETCO₂ inputs can be generated with "
        "assigned relationships, passed through predefined transfer functions, "
        "and combined into a CBFV output. MISO and SISO can then be compared "
        "with the same known pathways while recording duration, noise, input "
        "relationships, timing, and estimator settings are varied."
    ),
    (
        "This paper therefore has three connected objectives. The first is to "
        "develop and clearly document the MISO TFA pipeline and the reasoning "
        "behind its main choices. The second is to validate the pipeline using "
        "known-truth simulations and determine when MISO or SISO has lower "
        "pathway error. The third is to place cognitively normal recordings "
        "within the simulated operating space and determine whether the observed "
        "data support interpretable MISO estimates. The goal is to propose a "
        "reproducible reference workflow that can support future standardization. "
        "Comparisons between cognitively normal and mild cognitive impairment "
        "groups will be reserved for a separate paper after the method has been "
        "established."
    ),
]


REFERENCES = [
    "1. Claassen JAHR, Thijssen DHJ, Panerai RB, Faraci FM. Regulation of cerebral blood flow in humans: physiology and clinical implications of autoregulation. Physiological Reviews. 2021;101(4):1487-1559.",
    "2. Panerai RB, Brassard P, Burma JS, et al. Transfer function analysis of dynamic cerebral autoregulation: A CARNet white paper 2022 update. Journal of Cerebral Blood Flow & Metabolism. 2023;43(1):3-25.",
    "3. Hoiland RL, Fisher JA, Ainslie PN. Regulation of the cerebral circulation by arterial carbon dioxide. Comprehensive Physiology. 2019;9(3):1101-1154.",
    "4. Panerai RB, Simpson DM, Deverson ST, et al. Multivariate dynamic analysis of cerebral blood flow regulation in humans. IEEE Transactions on Biomedical Engineering. 2000;47(3):419-423.",
    "5. Peng T, Rowley AB, Ainslie PN, Poulin MJ, Payne SJ. Multivariate system identification for cerebral autoregulation. Annals of Biomedical Engineering. 2008;36(2):308-320.",
    "6. Kostoglou K, Bello-Robles F, Brassard P, et al. Time-domain methods for quantifying dynamic cerebral blood flow autoregulation: Review and recommendations. Journal of Cerebral Blood Flow & Metabolism. 2024;44(9):1480-1514.",
    "7. Marmarelis VZ, Shin DC, Orme ME, Zhang R. Model-based physiomarkers of cerebral hemodynamics in patients with mild cognitive impairment. Medical Engineering & Physics. 2014;36(5):628-637.",
    "8. Marmarelis VZ, Shin DC, Tarumi T, Zhang R. Comparison of model-based indices of cerebral autoregulation and vasomotor reactivity in patients with amnestic mild cognitive impairment. Journal of Alzheimer's Disease. 2017;56(1):89-105.",
    "9. Hashem S, Yamashiro S, Joe E, Chui H, Marmarelis V. PRBS gas challenges reveal impaired chemoreflex and cholinergic dynamics in MCI. Annals of Biomedical Engineering. 2026;54:2497-2507.",
]


def clear_paragraph_runs(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_paragraph(paragraph, text, template_run):
    run_properties = deepcopy(template_run._r.rPr) if template_run._r.rPr is not None else None
    clear_paragraph_runs(paragraph)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)
    return paragraph


def clone_paragraph_format(source, target):
    if source._p.pPr is not None:
        if target._p.pPr is not None:
            target._p.remove(target._p.pPr)
        target._p.insert(0, deepcopy(source._p.pPr))


def add_like(document, template_paragraph, text="", bold_lead=None):
    paragraph = document.add_paragraph()
    clone_paragraph_format(template_paragraph, paragraph)
    template_run = template_paragraph.runs[0] if template_paragraph.runs else None

    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        if template_run is not None and template_run._r.rPr is not None:
            lead._r.insert(0, deepcopy(template_run._r.rPr))
        lead.bold = True
        remainder = paragraph.add_run(text[len(bold_lead):])
        if template_run is not None and template_run._r.rPr is not None:
            remainder._r.insert(0, deepcopy(template_run._r.rPr))
    else:
        run = paragraph.add_run(text)
        if template_run is not None and template_run._r.rPr is not None:
            run._r.insert(0, deepcopy(template_run._r.rPr))
    return paragraph


def add_heading_like(document, heading_template, text):
    paragraph = add_like(document, heading_template, text)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_figure_space(document, body_template, figure_number, title, note):
    label = add_like(
        document,
        body_template,
        f"[Figure {figure_number} placeholder: {title}]",
    )
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label.runs:
        label.runs[0].bold = True

    note_paragraph = add_like(document, body_template, f"Figure note: {note}")
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note_paragraph.runs:
        note_paragraph.runs[0].italic = True

    reserved_space = add_like(document, body_template, "")
    reserved_space.paragraph_format.space_after = Pt(72)


def add_model_advantage_equation(document, body_template):
    paragraph = add_like(document, body_template, "")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    math_xml = f"""
    <m:oMathPara {nsdecls('m', 'w')}>
      <m:oMathParaPr><m:jc m:val="centerGroup"/></m:oMathParaPr>
      <m:oMath>
        <m:sSub>
          <m:sSubPr/>
          <m:e><m:r><m:t>A</m:t></m:r></m:e>
          <m:sub><m:r><m:t>model</m:t></m:r></m:sub>
        </m:sSub>
        <m:r><m:t> = </m:t></m:r>
        <m:sSub>
          <m:sSubPr/>
          <m:e><m:r><m:t>log</m:t></m:r></m:e>
          <m:sub><m:r><m:t>10</m:t></m:r></m:sub>
        </m:sSub>
        <m:r><m:t>(</m:t></m:r>
        <m:f>
          <m:fPr/>
          <m:num>
            <m:sSub>
              <m:sSubPr/>
              <m:e><m:r><m:t>E</m:t></m:r></m:e>
              <m:sub><m:r><m:t>SISO</m:t></m:r></m:sub>
            </m:sSub>
          </m:num>
          <m:den>
            <m:sSub>
              <m:sSubPr/>
              <m:e><m:r><m:t>E</m:t></m:r></m:e>
              <m:sub><m:r><m:t>MISO</m:t></m:r></m:sub>
            </m:sSub>
          </m:den>
        </m:f>
        <m:r><m:t>)</m:t></m:r>
      </m:oMath>
    </m:oMathPara>
    """
    paragraph._p.append(parse_xml(math_xml))


def main():
    document = Document(DOCUMENT)
    title_paragraph = document.paragraphs[0]
    heading_template = document.paragraphs[1]
    body_template = document.paragraphs[2]

    replace_paragraph(title_paragraph, TITLE, title_paragraph.runs[0])
    replace_paragraph(body_template, ABSTRACT, body_template.runs[0])

    add_heading_like(document, heading_template, "Introduction")
    for index, text in enumerate(INTRODUCTION_PARAGRAPHS, start=1):
        add_like(document, body_template, text)
        if index == 3:
            add_figure_space(
                document,
                body_template,
                1,
                "Comparison of SISO and MISO model structure",
                "Show two separate SISO models beside one joint MISO model. The figure should make clear that the MISO solution estimates the MAP and PETCO₂ transfer functions together while accounting for the relationship between the two inputs.",
            )
        if index == 4:
            add_figure_space(
                document,
                body_template,
                2,
                "Proposed MISO TFA analysis pipeline",
                "Show the flow from MAP, PETCO₂, and CBFV signals through preprocessing, Welch spectral estimation, the joint MISO solution, gain and phase calculation, conditioning and identifiability checks, statistical analysis, and organized figures and source-data exports.",
            )

    add_heading_like(document, heading_template, "Study Overview")
    add_like(
        document,
        body_template,
        "The study is organized into three connected parts. Together, these parts describe how the pipeline is built, how it is validated, and how it will be applied to real recordings.",
    )
    add_like(
        document,
        body_template,
        "Pipeline development. The first part defines the complete analysis workflow and explains why each major choice was made. This includes the required input signals, preprocessing, Welch spectral settings, the joint MISO solution, gain and phase calculation, conditioning diagnostics, regularization sensitivity, statistical testing, and the organization of outputs. The code is designed to remain modular and understandable so that another user can review or change each step.",
        bold_lead="Pipeline development.",
    )
    add_like(
        document,
        body_template,
        "Known-truth validation. The second part tests the pipeline against simulated MAP and PETCO₂ pathways that are known before estimation. Each simulated family represents one fixed physiological system. Recording duration, noise, signal alignment, and estimator settings can then be changed without changing that underlying truth. MISO and SISO will be compared using gain error, wrapped phase error, normalized complex error, and model advantage.",
        bold_lead="Known-truth validation.",
    )
    add_model_advantage_equation(document, body_template)
    add_like(
        document,
        body_template,
        "A positive model-advantage value means that MISO has lower error, while a negative value means that SISO has lower error. Statistical significance will be interpreted together with this direction. A significant difference alone will not be used to claim that MISO is better.",
    )
    add_like(
        document,
        body_template,
        "NC application. The third part places cognitively normal recordings within the simulated operating space. Recording duration, input coherence, spectral overlap, PETCO₂ fluctuation level, and conditioning will be evaluated before the MISO estimates are interpreted. Differences between MISO and SISO in NC data will be described as estimate differences rather than known-truth accuracy differences.",
        bold_lead="NC application.",
    )

    add_heading_like(document, heading_template, "Selected References")
    for reference in REFERENCES:
        paragraph = add_like(document, body_template, reference)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)

    document.save(DOCUMENT)


if __name__ == "__main__":
    main()
