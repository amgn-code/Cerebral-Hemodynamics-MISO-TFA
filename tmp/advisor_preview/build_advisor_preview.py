from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name(
    "MISO Approach Paper - Working Abstract and Introduction.docx"
)

NAVY = "183A5A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66717D"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D7DCE2"
WHITE = "FFFFFF"
BLACK = "000000"


def set_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.extend([fonts, size, color])
    text = OxmlElement("w:t")
    text.text = "1"
    r.extend([r_pr, text])
    field.append(r)
    paragraph._p.append(field)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Caption",):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(9.5)
        style.font.italic = True
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.1

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("MISO approach paper | Directional manuscript preview")
    set_font(hr, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_number(fp)


def add_body_paragraph(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25) if first_line else None
    p.paragraph_format.keep_together = False
    r = p.add_run(text)
    set_font(r)
    return p


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.keep_together = True
    label_run = p.add_run(label + " ")
    set_font(label_run, bold=True, color=DARK_BLUE)
    body_run = p.add_run(text)
    set_font(body_run)
    return p


def add_figure_placeholder(doc, number, title, note):
    table = doc.add_table(rows=1, cols=1)
    keep_table_row_together(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color="9CB3C7", size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.height = Inches(1.20)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    r = p.add_run(f"FIGURE {number} PLACEHOLDER")
    set_font(r, size=11, color=NAVY, bold=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(5)
    r2 = p2.add_run(title)
    set_font(r2, size=11, color=BLACK, bold=True)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.15
    r3 = p3.add_run("Figure note: " + note)
    set_font(r3, size=9.5, color=MUTED, italic=True)

    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = False
    caption.add_run(f"Figure {number}. {title}. This marker will be replaced after the final analysis and figure design are complete.")


def add_key_question(doc):
    table = doc.add_table(rows=1, cols=1)
    keep_table_row_together(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color="B8C3CE", size=6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Central question")
    set_font(r, size=10, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(
        "When does adding PETCO₂ to a cerebral-hemodynamic transfer function improve recovery of the underlying MAP and PETCO₂ pathways, and when does the additional model complexity become unsupported by the available recording?"
    )
    set_font(r2, size=10.5, color=BLACK)


def add_roadmap_table(doc):
    rows = [
        ("Generator validation", "Did the simulation create the assigned input relationships and true pathway properties?"),
        ("Reference-condition accuracy", "Under the prespecified 300-second condition, which model is closer to the known transfer functions?"),
        ("Recording duration and noise", "How do finite recording length and separate noise sources change SISO and MISO error?"),
        ("Input identifiability", "How do frequency-dependent input coherence, PSD-shape overlap, and conditioning affect pathway separation?"),
        ("PETCO₂ excitation and contribution", "When do PETCO₂ fluctuations and its contribution to CBFV provide enough information to justify the second pathway?"),
        ("Delay, synchronization, and estimator settings", "How sensitive are estimates to physiological delay, timing error, Welch design, and regularization?"),
        ("NC operating-space placement", "Do the cognitively normal recordings fall within simulation conditions that support interpretable MISO estimation?"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [3000, 6360])
    set_table_borders(table, color=MID_GRAY, size=6)
    hdr = table.rows[0].cells
    for idx, text in enumerate(("Planned evidence", "Question answered")):
        set_cell_shading(hdr[idx], "E8EEF5")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=10, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    keep_table_row_together(table.rows[0])

    for left, right in rows:
        cells = table.add_row().cells
        keep_table_row_together(table.rows[-1])
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
        lp = cells[0].paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(left)
        set_font(lr, size=9.5, color=DARK_BLUE, bold=True)
        rp = cells[1].paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(right)
        set_font(rr, size=9.5)

    set_table_geometry(table, [3000, 6360])


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(f"{number}. {text}")
    set_font(r, size=9.5)


def build_document():
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "MISO Approach Paper: Working Abstract and Introduction"
    doc.core_properties.subject = "Directional manuscript preview for advisor discussion"
    doc.core_properties.author = "Amogh Nair and collaborators"
    doc.core_properties.comments = "Working draft based on the Paper 1 known-truth simulation framework"

    # Compact academic adaptation of the editorial-cover pattern.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("WORKING ABSTRACT AND INTRODUCTION")
    set_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(
        "Evaluating Two-Input Transfer Function Analysis of Cerebral Hemodynamics"
    )
    set_font(r, size=24, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(
        "A known-truth simulation framework and application to cognitively normal data"
    )
    set_font(r, size=13, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Directional manuscript preview for discussion with Sho")
    set_font(r, size=10.5, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("August 3, 2026 | Author names and affiliations to be finalized")
    set_font(r, size=9.5, color=MUTED)

    add_key_question(doc)

    doc.add_heading("Abstract", level=1)
    abstract = (
        "Transfer function analysis of cerebral hemodynamics commonly treats mean arterial pressure (MAP) as a single input and cerebral blood flow velocity (CBFV) as the output. CBFV also responds to arterial carbon dioxide, which is commonly represented during resting recordings by end-tidal carbon dioxide (PETCO₂). When MAP and PETCO₂ share frequency content, a MAP-only single-input, single-output (SISO) transfer function can contain variation associated with both physiological drivers. A multiple-input, single-output (MISO) model can estimate the MAP-to-CBFV and PETCO₂-to-CBFV relationships simultaneously, but the second input does not guarantee a better estimate. Short recordings, weak PETCO₂ fluctuations, correlated inputs, measurement noise, timing errors, and poorly conditioned spectral matrices may offset the conceptual benefit of separating the two pathways. This study therefore uses a known-truth simulation framework to identify the conditions under which two-input MISO transfer function analysis recovers the true MAP and PETCO₂ pathways more accurately than separate SISO models. Each simulated family represents a fixed physiological system, while recording duration, noise, timing alignment, and estimator settings are varied as observation conditions without changing the underlying truth. Accuracy will be evaluated using frequency-resolved gain error, wrapped phase error, normalized complex error, and a signed model-advantage measure that indicates whether MISO or SISO has lower error. Resting recordings from cognitively normal adults will then be placed within the simulated operating space to determine whether their duration, input relationships, excitation, and conditioning support interpretable MISO estimates. The purpose of this first paper is not to claim that MISO is universally superior or to test a disease biomarker. Its purpose is to establish when the additional input is justified, when conventional SISO remains sufficient, and when neither separated pathway should be interpreted without additional data or stabilization. This foundation will support a later study of MISO-derived markers in mild cognitive impairment."
    )
    add_body_paragraph(doc, abstract, first_line=False)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "Keywords: cerebral autoregulation; carbon dioxide; end-tidal carbon dioxide; transcranial Doppler; transfer function analysis; multiple-input system; simulation; system identification"
    )
    set_font(r, size=9.5, color=MUTED, italic=True)

    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Why cerebral-hemodynamic transfer functions are useful", level=2)
    add_body_paragraph(
        doc,
        "The brain requires a continuous blood supply even though arterial pressure changes from moment to moment. Cerebral autoregulation describes the physiological processes that limit how strongly these pressure changes alter cerebral perfusion [1]. Dynamic cerebral autoregulation focuses on responses that can be observed over seconds to minutes. In a common resting protocol, MAP is treated as an input, transcranial Doppler CBFV is treated as an output, and transfer function analysis describes their relationship at each frequency [2].",
    )
    add_body_paragraph(
        doc,
        "A transfer function is complex-valued. Its magnitude is reported as gain, while its angle is reported as phase. Gain describes how strongly an input fluctuation is represented in the output at a given frequency. Phase describes the timing relationship between the input and output on the frequency-domain circle. Coherence describes the consistency of their linear relationship at that frequency. Together, these measures retain information that would be lost if a recording were reduced to one whole-range correlation or one average response. They do not prove causation by themselves, but they provide an interpretable description of cerebral hemodynamic dynamics.",
    )

    doc.add_heading("1.2 The limitation of treating MAP as the only input", level=2)
    add_body_paragraph(
        doc,
        "MAP is not the only physiological input that can influence CBFV. Carbon dioxide has a strong effect on cerebrovascular tone, and changes in arterial carbon dioxide can alter cerebral blood flow over frequencies that overlap those used for autoregulation analysis [3]. PETCO₂ is not identical to arterial carbon dioxide, but it provides a practical continuous estimate during resting recordings. A model that omits PETCO₂ may therefore assign some carbon-dioxide-associated variation to the pressure pathway.",
    )
    add_body_paragraph(
        doc,
        "If MAP and PETCO₂ were unrelated, a MAP-only model and a two-input model would be expected to provide similar MAP pathway estimates, apart from ordinary finite-record variation. If the inputs are related, however, the MAP-to-CBFV cross-spectrum can contain both the true MAP-associated response and a PETCO₂-associated response carried through the MAP-to-PETCO₂ relationship. The same problem applies in reverse to a PETCO₂-only model. Separate SISO estimates therefore describe pairwise associations. They should not automatically be interpreted as pathways that have been separated from the other measured input.",
    )
    add_body_paragraph(
        doc,
        "This concern has a clear mathematical basis. Under a linear two-input system, the MAP-only SISO estimate contains the true MAP pathway plus a term that depends on the PETCO₂ pathway and the MAP-PETCO₂ cross-spectrum. That additional term becomes zero only when the inputs are uncorrelated at that frequency or when the omitted PETCO₂ pathway contributes nothing. The direction and size of the distortion can therefore change across frequency rather than acting as one constant bias.",
    )

    add_figure_placeholder(
        doc,
        1,
        "Conceptual distinction between SISO and MISO cerebral-hemodynamic models",
        "Use two matched panels. The SISO panel should show MAP-to-CBFV and PETCO₂-to-CBFV estimated separately, with the omitted pathway visibly capable of entering the pairwise estimate. The MISO panel should show MAP and PETCO₂ entering one joint solve so that each estimated transfer function is conditional on the other measured input. Keep the diagram conceptual and avoid implying that the model proves biological causation.",
    )

    doc.add_heading("1.3 Existing work supports multiple inputs but does not settle the practical question", level=2)
    add_body_paragraph(
        doc,
        "The two-input idea is not new. Panerai and colleagues modeled MAP and end-tidal carbon dioxide together in multivariable analysis of cerebral blood flow regulation [4]. Peng and colleagues developed multivariable system-identification methods for cerebral autoregulation and showed that gas-related inputs could provide additional low-frequency information [5]. Marmarelis and colleagues have also used pressure and PETCO₂ as joint inputs when characterizing cerebral hemodynamics and developing model-based physiomarkers [7,8]. More recent work using broadband gas challenges further shows that multivariable cerebral-hemodynamic modeling remains an active area of development [9].",
    )
    add_body_paragraph(
        doc,
        "The CARNet recommendations provide an important framework for TFA design, reporting, and interpretation [2], while the CARNet time-domain white paper identifies multivariable models as an important extension of conventional analysis [6]. Together, these studies establish that MAP and carbon dioxide can reasonably be treated as joint inputs. They do not, however, remove the need to determine when a two-input frequency-domain estimator is more accurate than the conventional one-input approach under realistic resting-recording conditions.",
    )
    add_body_paragraph(
        doc,
        "This distinction is central to the proposed paper. Its novelty is not the claim that PETCO₂ can influence CBFV or that a MISO equation can be solved. The contribution is a transparent known-truth benchmark that measures both sides of the tradeoff: the benefit of separating a shared-input contribution and the cost of estimating two pathway coefficients from finite, noisy, and potentially redundant data.",
    )

    doc.add_heading("1.4 Why MISO is not automatically better", level=2)
    add_body_paragraph(
        doc,
        "Adding PETCO₂ gives the model a way to distinguish pressure-associated and carbon-dioxide-associated variation, but it also creates an identifiability problem. At each frequency, MISO must estimate two pathway coefficients from the same finite recording. If MAP and PETCO₂ become nearly indistinguishable, the input spectral matrix becomes poorly conditioned. Small spectral-estimation errors can then produce large changes in the separated transfer functions. The problem can be amplified by short recordings, low input power, measurement noise, timing error, and estimator choices.",
    )
    add_body_paragraph(
        doc,
        "Input coherence is therefore not simply good or bad. With nearly independent inputs, conventional SISO may already have little shared-input contamination, so MISO can offer limited benefit while still estimating an extra coefficient. With moderate coupling, SISO mixing may become important while enough independent variation remains for MISO to separate the pathways. With coherence near one, the normalized determinant of the input spectral matrix approaches zero and the MISO solution becomes sensitive to small errors. The useful region may therefore be intermediate rather than increasing monotonically with coherence.",
    )
    add_body_paragraph(
        doc,
        "Spectral overlap provides related but different information. Two inputs may occupy similar frequency regions without being strongly coherent, and coherence is difficult to estimate where either input has little power. The proposed analysis therefore retains frequency-dependent coherence, input power, PSD-shape overlap, and conditioning as separate diagnostics. This is intended to make the model's limitations visible rather than hiding them behind one average coherence value.",
    )

    doc.add_heading("1.5 Why known-truth simulation is necessary", level=2)
    add_body_paragraph(
        doc,
        "Human recordings cannot reveal the true MAP-to-CBFV or PETCO₂-to-CBFV transfer functions. If SISO and MISO estimates differ in an empirical recording, the difference alone does not show which estimate is closer to the underlying physiology. A higher in-sample multiple coherence also does not prove more accurate pathway separation. The model with more inputs will often explain at least as much of the observed output even when its individual pathway estimates are unstable.",
    )
    add_body_paragraph(
        doc,
        "Known-truth simulation makes the accuracy question testable. The simulation generates MAP and PETCO₂ inputs with assigned spectral properties, passes them through predefined MAP and PETCO₂ transfer functions, and combines their separate contributions into a simulated CBFV output. SISO and MISO can then be applied to exactly the same observations and compared with the known complex transfer functions. Because truth is available, the analysis can distinguish a visually smooth estimate from an accurate estimate.",
    )
    add_body_paragraph(
        doc,
        "The simulation is organized hierarchically. A family represents one fixed physiological system. Its input relationships, true pathway shapes, relative pathway contribution, and physiological delay remain unchanged. Multiple observations can then be created from the same family by changing how it is recorded or analyzed, including duration, output noise, input measurement noise, timing alignment, Welch settings, and regularization. This separation permits paired comparisons because each family is evaluated under the same controlled observation levels.",
    )

    add_figure_placeholder(
        doc,
        2,
        "Known-truth family, observation, estimation, and evaluation workflow",
        "Show the pipeline from a fixed simulated family to repeated observation conditions. The family block should include input spectra, frequency-dependent MAP-PETCO₂ coherence, true MAP and PETCO₂ pathways, pathway contribution, and physiological delay. The observation block should include duration, CBFV noise, MAP noise, PETCO₂ noise, and timing alignment. The final block should show SISO and MISO estimates compared with truth using gain, wrapped phase, complex error, and signed model advantage.",
    )

    doc.add_heading("1.6 The role of the cognitively normal recordings", level=2)
    add_body_paragraph(
        doc,
        "The first paper will use cognitively normal recordings for an empirical application, but the NC data serve a different purpose from the known-truth simulation. The simulation determines estimator behavior because the true pathways are known. The NC analysis determines whether the available human recordings occupy conditions represented by the simulation. Recording duration, input coherence, PSD overlap, PETCO₂ fluctuation amplitude, and conditioning will be compared with the simulated operating space before the MISO estimates are interpreted.",
    )
    add_body_paragraph(
        doc,
        "This distinction limits overinterpretation. In simulation, model advantage compares each estimate with known truth. In NC data, only the two estimates are observed. The simulation can support the interpretation of an NC MISO-versus-SISO difference when the empirical operating point is covered, but it cannot assign an exact unobserved error to an individual participant. NC-versus-MCI comparisons are therefore reserved for a second paper, after the estimator and its operating boundaries have been established.",
    )

    doc.add_heading("1.7 Study objectives and hypotheses", level=2)
    add_labeled_paragraph(
        doc,
        "Objective 1.",
        "Develop and validate a transparent two-input signal generator in which realized input and pathway properties can be compared with their assigned values.",
    )
    add_labeled_paragraph(
        doc,
        "Objective 2.",
        "Compare SISO and MISO gain, phase, and complex transfer-function errors under a prespecified 300-second reference observation.",
    )
    add_labeled_paragraph(
        doc,
        "Objective 3.",
        "Determine how recording duration, separate noise sources, input coherence, PSD overlap, PETCO₂ excitation, pathway contribution, true delay, timing misalignment, and estimator settings alter model performance.",
    )
    add_labeled_paragraph(
        doc,
        "Objective 4.",
        "Place NC recordings within the simulation-defined operating space and describe NC MISO-versus-SISO differences without making disease-marker claims.",
    )
    add_body_paragraph(
        doc,
        "The primary hypothesis is conditional. MISO is expected to provide the greatest benefit when PETCO₂ makes a meaningful contribution to CBFV and the two measured inputs share enough variation for SISO mixing to matter, while retaining enough independent information for the pathways to be separated. Its advantage is expected to decrease with short records, weak PETCO₂ excitation, high measurement noise, timing error, and severe input collinearity. The boundaries of this useful region will be treated as empirical results rather than selected visually or assumed in advance.",
    )

    doc.add_heading("2. Proposed approach at a glance", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "This section is included for advisor discussion. It summarizes how the research question will be answered without reproducing the full Methods section of the manuscript."
    )
    set_font(r, size=10, color=MUTED, italic=True)

    doc.add_heading("2.1 Simulation structure", level=2)
    add_labeled_paragraph(
        doc,
        "Family-level properties.",
        "Input coherence structure, PSD-shape similarity, the PETCO₂-to-MAP fluctuation SD ratio, the PETCO₂-to-MAP pathway band-gain ratio, true PETCO₂ delay, and pathway time constants define the underlying simulated system.",
    )
    add_labeled_paragraph(
        doc,
        "Observation-level properties.",
        "Recording duration, CBFV output noise, MAP measurement noise, PETCO₂ measurement noise, and PETCO₂ timing misalignment describe how the fixed system is observed.",
    )
    add_labeled_paragraph(
        doc,
        "Estimator-level properties.",
        "Welch window length, Welch overlap, spectral smoothing, and scale-normalized ridge regularization describe how each observation is analyzed.",
    )
    add_body_paragraph(
        doc,
        "The planned paper-profile simulation will use 1,000 families. The reference observation is 300 seconds long, sampled at 4 Hz, and analyzed from 0.005 to 0.35 Hz using 128-second Hann windows with 50% overlap. Additional durations extend from the lower feasible range to 896 seconds. Output and input noise are varied from clean conditions through 0 dB, where signal and noise have equal power. The broad design is intended to reveal operating boundaries rather than imitate only the current NC cohort.",
    )

    doc.add_heading("2.2 Accuracy and model advantage", level=2)
    add_body_paragraph(
        doc,
        "Gain error and wrapped phase error retain the two familiar components of the transfer function. Normalized complex error provides a complementary whole-transfer-function measure by calculating the squared complex-plane distance between the estimate and truth and normalizing it to the true pathway energy. It is a real, nonnegative error even though it is calculated from complex-valued transfer functions.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for text, subscript in (
        ("A", False), ("model", True), (" = log", False), ("10", True),
        ("(", False), ("E", False), ("SISO", True), (" / ", False),
        ("E", False), ("MISO", True), (")", False),
    ):
        r = p.add_run(text)
        set_font(r, name="Cambria Math", size=12, color=NAVY, bold=True)
        r.font.subscript = subscript
    add_body_paragraph(
        doc,
        "Positive model advantage means MISO has lower error. Negative model advantage means SISO has lower error. A value of zero means the errors are equal. For example, a value of 1 means SISO error is ten times MISO error, while a value of -1 means MISO error is ten times SISO error. The current heatmap display range of -2 to 2 is a plotting limit, not a normalization rule.",
    )

    doc.add_heading("2.3 Statistical interpretation", level=2)
    add_body_paragraph(
        doc,
        "The simulated family is the independent statistical unit. Frequency bins, Welch windows, and repeated observations from the same family will not be treated as independent samples. At each prespecified condition or heatmap cell, family-level model advantage will be tested against zero and reported with its mean, standard deviation, 95% confidence interval, probability value, adjusted probability value when applicable, and valid family count.",
    )
    add_body_paragraph(
        doc,
        "Statistical significance will always be interpreted together with direction. A significant positive advantage supports lower MISO error. A significant negative advantage supports lower SISO error. A nonsignificant result does not establish equivalence. Benjamini-Hochberg adjustment will be applied within defined families of frequency-bin, factor-level, or heatmap-cell comparisons [11]. In addition, significance at 896 seconds but not at 300 seconds will not be described as proof that duration changed model advantage. That claim requires a direct paired contrast or repeated-measures model comparing the two durations.",
    )

    doc.add_heading("2.4 NC application", level=2)
    add_body_paragraph(
        doc,
        "The NC analysis will first assess simulation coverage. Each recording will be characterized by duration, frequency-dependent input coherence, PSD-shape overlap, PETCO₂-to-MAP fluctuation ratio, conditioning, and other prespecified diagnostics. Values outside the simulated range will be identified rather than silently extrapolated. If an important part of the NC distribution is not covered, the simulation should be expanded before it is used to justify the empirical estimator.",
    )
    add_body_paragraph(
        doc,
        "Only after coverage is established will NC MISO and SISO transfer-function estimates be compared. These empirical differences will be reported as estimate differences, not known-truth model advantage. Gain and wrapped phase will be shown across frequency and prespecified bands, with input power, coherence, and conditioning presented nearby so that apparently large model differences can be interpreted in the context of data support.",
    )

    add_figure_placeholder(
        doc,
        3,
        "Placement of cognitively normal recordings within the simulated operating space",
        "Show simulation coverage and the locations of NC recordings for duration, frequency-dependent input coherence, PSD overlap, PETCO₂ fluctuation ratio, and conditioning. Distinguish well-covered regions from sparse or unsupported regions. This figure should justify where the simulation can inform NC interpretation and identify where additional simulation is required.",
    )

    doc.add_heading("3. Expected contribution and claim boundaries", level=1)
    add_body_paragraph(
        doc,
        "The proposed paper is intended to establish a practical foundation for two-input cerebral TFA. Its main product should be a map of the conditions under which MISO has lower known-truth error, the conditions under which SISO is sufficient, and the conditions under which neither separated pathway should be trusted without more data or stabilization. The simulation should also produce reporting recommendations for recording duration, input excitation, coherence, PSD overlap, conditioning, timing, Welch design, and regularization.",
    )
    add_labeled_paragraph(
        doc,
        "The paper can claim:",
        "known-truth differences in estimator accuracy under simulated conditions, the direction and uncertainty of those differences, and whether the available NC recordings fall within the studied operating space.",
    )
    add_labeled_paragraph(
        doc,
        "The paper should not claim:",
        "that MISO is universally superior, that coherence alone proves model validity, that quick-simulation output is final evidence, that an empirical NC model difference proves greater accuracy, or that NC-versus-MCI biomarkers have already been established.",
    )
    add_labeled_paragraph(
        doc,
        "Paper 2 remains separate:",
        "once the MISO estimator and its operating requirements have been established, the later paper can ask whether conditional MAP and PETCO₂ pathway metrics reveal reproducible biological differences between NC and MCI participants.",
    )

    doc.add_heading("4. Planned evidence sequence", level=1)
    add_body_paragraph(
        doc,
        "The final manuscript can contain many detailed plots, but the main narrative should move through a small number of evidence groups. The table below summarizes the role of each group.",
    )
    add_roadmap_table(doc)

    doc.add_heading("References", level=1)
    references = [
        "Claassen JAHR, Thijssen DHJ, Panerai RB, Faraci FM. Regulation of cerebral blood flow in humans: physiology and clinical implications of autoregulation. Physiological Reviews. 2021;101(4):1487-1559. doi:10.1152/physrev.00022.2020.",
        "Panerai RB, Brassard P, Burma JS, et al.; Cerebrovascular Research Network. Transfer function analysis of dynamic cerebral autoregulation: A CARNet white paper 2022 update. Journal of Cerebral Blood Flow & Metabolism. 2023;43(1):3-25. doi:10.1177/0271678X221119760.",
        "Hoiland RL, Fisher JA, Ainslie PN. Regulation of the cerebral circulation by arterial carbon dioxide. Comprehensive Physiology. 2019;9(3):1101-1154. doi:10.1002/cphy.c180021.",
        "Panerai RB, Simpson DM, Deverson ST, Mahony P, Hayes P, Evans DH. Multivariate dynamic analysis of cerebral blood flow regulation in humans. IEEE Transactions on Biomedical Engineering. 2000;47(3):419-423. doi:10.1109/10.827312.",
        "Peng T, Rowley AB, Ainslie PN, Poulin MJ, Payne SJ. Multivariate system identification for cerebral autoregulation. Annals of Biomedical Engineering. 2008;36(2):308-320. doi:10.1007/s10439-007-9412-9.",
        "Kostoglou K, Bello-Robles F, Brassard P, et al. Time-domain methods for quantifying dynamic cerebral blood flow autoregulation: Review and recommendations. A white paper from the Cerebrovascular Research Network. Journal of Cerebral Blood Flow & Metabolism. 2024;44(9):1480-1514. doi:10.1177/0271678X241249276.",
        "Marmarelis VZ, Shin DC, Orme ME, Zhang R. Model-based physiomarkers of cerebral hemodynamics in patients with mild cognitive impairment. Medical Engineering & Physics. 2014;36(5):628-637. doi:10.1016/j.medengphy.2014.02.025.",
        "Marmarelis VZ, Shin DC, Tarumi T, Zhang R. Comparison of model-based indices of cerebral autoregulation and vasomotor reactivity using transcranial Doppler versus near-infrared spectroscopy in patients with amnestic mild cognitive impairment. Journal of Alzheimer's Disease. 2017;56(1):89-105. doi:10.3233/JAD-161004.",
        "Hashem S, Yamashiro S, Joe E, Chui H, Marmarelis V. PRBS gas challenges reveal impaired chemoreflex and cholinergic dynamics in MCI. Annals of Biomedical Engineering. 2026;54:2497-2507. doi:10.1007/s10439-026-04213-7.",
        "Welch PD. The use of fast Fourier transform for the estimation of power spectra: A method based on time averaging over short, modified periodograms. IEEE Transactions on Audio and Electroacoustics. 1967;15(2):70-73. doi:10.1109/TAU.1967.1161901.",
        "Benjamini Y, Hochberg Y. Controlling the false discovery rate: A practical and powerful approach to multiple testing. Journal of the Royal Statistical Society: Series B. 1995;57(1):289-300. doi:10.1111/j.2517-6161.1995.tb02031.x.",
    ]
    for number, reference in enumerate(references, start=1):
        add_reference(doc, number, reference)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Working document for discussion. Final wording, journal format, figures, numerical results, and author information remain to be completed."
    )
    set_font(r, size=9, color=MUTED, italic=True)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
