from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCUMENT_PATH = Path(__file__).with_name("Abstract & Introduction.docx")


def find_paragraph(document, beginning):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(beginning)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning with {beginning!r}; "
            f"found {len(matches)}."
        )
    return matches[0]


def replace_plain_paragraph(paragraph, text):
    """Replace body text while preserving paragraph formatting."""
    paragraph.clear()
    paragraph.add_run(text)


def replace_text_before_first_hyperlink(paragraph, text):
    """Replace introductory text without disturbing the linked title."""
    hyperlink = next(
        child
        for child in paragraph._p
        if child.tag == qn("w:hyperlink")
    )
    first_run = next(
        child
        for child in paragraph._p
        if child.tag == qn("w:r")
    )

    for child in list(paragraph._p):
        if child is hyperlink:
            break
        if child.tag == qn("w:r") and child is not first_run:
            paragraph._p.remove(child)

    for child in list(first_run):
        if child.tag != qn("w:rPr"):
            first_run.remove(child)
    text_element = first_run.makeelement(qn("w:t"))
    text_element.text = text
    first_run.append(text_element)


document = Document(DOCUMENT_PATH)
hyperlinks_before = len(document.element.body.findall(".//w:hyperlink", document.element.nsmap))

abstract = find_paragraph(document, "Transfer Function Analysis (TFA)")
abstract_text = abstract.text
abstract_text = abstract_text.replace(
    "Transfer functions estimated from one simulated realization will also "
    "be applied to an independent realization to calculate normalized CBFV "
    "prediction error and determine whether improved prediction is "
    "accompanied by more accurate pathway recovery.",
    "Transfer functions estimated from one simulated recording will also be "
    "tested on a second independently generated recording from the same "
    "underlying system. This allows CBFV prediction on new data to be compared "
    "with the accuracy of the recovered MAP and PETCO₂ pathways.",
)
abstract_text = abstract_text.replace(
    "Resting baseline recordings from cognitively normal (NC) adults will then "
    "be placed within the simulated operating space to determine whether their "
    "duration, input relationships, excitation, and conditioning support "
    "interpretable MISO estimates.",
    "Properties of resting baseline recordings from cognitively normal (NC) "
    "adults, including recording duration, the relationship between MAP and "
    "PETCO₂, PETCO₂ fluctuation size, and the numerical stability of the "
    "MISO solution, will then be compared with the simulation results. This "
    "will show whether the recordings resemble simulated conditions in which "
    "MISO recovered the known pathways reliably.",
)
if "realization" in abstract_text or "simulated operating space" in abstract_text:
    raise RuntimeError("The abstract clarity replacements were incomplete.")
replace_plain_paragraph(abstract, abstract_text)

panerai_limit = find_paragraph(
    document,
    "However, a reduction in CBFV prediction error",
)
replace_text_before_first_hyperlink(
    panerai_limit,
    "However, lower CBFV prediction error does not necessarily mean that the "
    "separate MAP and PETCO₂ transfer functions were estimated correctly. "
    "MAP, PETCO₂, and CBFV can be measured in human recordings, but the true "
    "MAP-to-CBFV and PETCO₂-to-CBFV pathways are not independently known. "
    "Panerai et al. therefore showed that adding PETCO₂ improved the "
    "description of total CBFV, but they could not directly compare the "
    "estimated pathways with known physiological pathways. The authors also "
    "reported no significant time-domain cross-correlation between MAP and "
    "PETCO₂. This does not describe the relationship between the inputs at "
    "every frequency, so the accuracy of pathway separation under high "
    "frequency-specific coherence or poor spectral conditioning remains "
    "uncertain. Peng et al. also supported the use of multivariable system "
    "identification for cerebral autoregulation (",
)

simulation = find_paragraph(
    document,
    "Subject recordings cannot show the true MAP-to-CBFV",
)
replace_plain_paragraph(
    simulation,
    "Subject recordings cannot show the true MAP-to-CBFV or PETCO₂-to-CBFV "
    "transfer functions. If MISO and SISO produce different results in an "
    "observed recording, that difference alone cannot show which model is "
    "closer to the underlying physiology. A known-truth simulation makes this "
    "comparison possible. MAP and PETCO₂ inputs can be generated with "
    "assigned relationships, passed through predefined transfer functions, "
    "and combined into a CBFV output. Transfer functions estimated from one "
    "simulated recording can be compared directly with the known pathways and "
    "then tested on a second independently generated recording from the same "
    "simulated physiological system. The second recording has the same true "
    "pathways and simulation settings but different randomly generated signal "
    "fluctuations and noise. This allows pathway accuracy and CBFV prediction "
    "on new data to be evaluated together while recording duration, noise, "
    "input relationships, timing, and estimator settings are varied.",
)

objectives = find_paragraph(
    document,
    "This paper therefore has three connected objectives.",
)
replace_plain_paragraph(
    objectives,
    "This paper therefore has three connected objectives. The first is to "
    "develop and clearly document the MISO TFA pipeline and the reasoning "
    "behind its main choices. The second is to validate the pipeline using "
    "known-truth simulations and determine when MISO or SISO has lower pathway "
    "error. The third is to compare measurable properties of the NC recordings "
    "with the simulation results and determine whether the recordings resemble "
    "conditions in which MISO recovered the known pathways reliably. The goal "
    "is to propose a reproducible reference workflow that can support future "
    "research.",
)

hyperlinks_after = len(document.element.body.findall(".//w:hyperlink", document.element.nsmap))
if hyperlinks_after != hyperlinks_before:
    raise RuntimeError(
        f"Hyperlink count changed from {hyperlinks_before} to {hyperlinks_after}."
    )

document.save(DOCUMENT_PATH)
print(f"Updated {DOCUMENT_PATH}")
print(f"Hyperlinks preserved: {hyperlinks_after}")
