from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


DOCUMENT_PATH = "Abstract & Introduction.docx"


def clear_paragraph(paragraph):
    """Remove paragraph contents while keeping its paragraph-level formatting."""
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def delete_paragraph(paragraph):
    paragraph_element = paragraph._element
    paragraph_element.getparent().remove(paragraph_element)


def copy_run_style(source_run, target_run):
    if source_run is not None and source_run._element.rPr is not None:
        target_run._element.insert(0, deepcopy(source_run._element.rPr))


def add_hyperlink(paragraph, text, url, source_run=None):
    """Add a visible, clickable hyperlink while retaining the body font."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    if source_run is not None and source_run._element.rPr is not None:
        run_properties = deepcopy(source_run._element.rPr)
    else:
        run_properties = OxmlElement("w:rPr")

    color = run_properties.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        run_properties.append(color)
    color.set(qn("w:val"), "0563C1")

    underline = run_properties.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        run_properties.append(underline)
    underline.set(qn("w:val"), "single")

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_with_linked_parts(paragraph, parts):
    """Replace a paragraph using (text, optional_url) parts."""
    source_run = paragraph.runs[0] if paragraph.runs else None
    source_run_properties = (
        deepcopy(source_run._element.rPr)
        if source_run is not None and source_run._element.rPr is not None
        else None
    )
    clear_paragraph(paragraph)

    for text, url in parts:
        if url:
            add_hyperlink(paragraph, text, url, source_run)
        else:
            run = paragraph.add_run(text)
            if source_run_properties is not None:
                run._element.insert(0, deepcopy(source_run_properties))


def set_plain_heading(paragraph, text, template_run=None, size=12):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    copy_run_style(template_run, run)
    run.bold = True
    run.italic = False
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None


def add_reference_group(document, body_template, heading, purpose, references):
    heading_paragraph = document.add_paragraph()
    set_plain_heading(heading_paragraph, heading, body_template.runs[0], size=12)

    purpose_paragraph = document.add_paragraph()
    purpose_run = purpose_paragraph.add_run("Use in the papers: " + purpose)
    copy_run_style(body_template.runs[0], purpose_run)
    purpose_run.italic = True
    purpose_paragraph.paragraph_format.space_after = Pt(5)

    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)

        author_separator = " " if reference["authors"].endswith(".") else ". "
        prefix_run = paragraph.add_run(reference["authors"] + author_separator)
        copy_run_style(body_template.runs[0], prefix_run)
        add_hyperlink(
            paragraph,
            reference["title"],
            reference["url"],
            body_template.runs[0],
        )
        suffix_run = paragraph.add_run(". " + reference["citation"])
        copy_run_style(body_template.runs[0], suffix_run)


references = {
    "foundations": [
        {
            "authors": "Claassen JAHR, Thijssen DHJ, Panerai RB, Faraci FM",
            "title": "Regulation of cerebral blood flow in humans: physiology and clinical implications of autoregulation",
            "citation": "Physiological Reviews. 2021;101(4):1487-1559. doi:10.1152/physrev.00022.2020",
            "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8576366/",
        },
        {
            "authors": "Panerai RB, Brassard P, Burma JS, et al.; Cerebrovascular Research Network",
            "title": "Transfer function analysis of dynamic cerebral autoregulation: A CARNet white paper 2022 update",
            "citation": "Journal of Cerebral Blood Flow & Metabolism. 2023;43(1):3-25. doi:10.1177/0271678X221119760",
            "url": "https://doi.org/10.1177/0271678X221119760",
        },
        {
            "authors": "Kostoglou K, Bello-Robles F, Brassard P, et al.",
            "title": "Time-domain methods for quantifying dynamic cerebral blood flow autoregulation: Review and recommendations. A white paper from the Cerebrovascular Research Network (CARNet)",
            "citation": "Journal of Cerebral Blood Flow & Metabolism. 2024;44(9):1480-1514. doi:10.1177/0271678X241249276",
            "url": "https://journals.sagepub.com/doi/pdf/10.1177/0271678X241249276",
        },
        {
            "authors": "Aaslid R, Markwalder TM, Nornes H",
            "title": "Noninvasive transcranial Doppler ultrasound recording of flow velocity in basal cerebral arteries",
            "citation": "Journal of Neurosurgery. 1982;57(6):769-774. doi:10.3171/jns.1982.57.6.0769",
            "url": "https://doi.org/10.3171/jns.1982.57.6.0769",
        },
    ],
    "multivariable": [
        {
            "authors": "Hoiland RL, Fisher JA, Ainslie PN",
            "title": "Regulation of the cerebral circulation by arterial carbon dioxide",
            "citation": "Comprehensive Physiology. 2019;9(3):1101-1154. doi:10.1002/cphy.c180021",
            "url": "https://pubmed.ncbi.nlm.nih.gov/31187899/",
        },
        {
            "authors": "Panerai RB, Simpson DM, Deverson ST, Mahony P, Hayes P, Evans DH",
            "title": "Multivariate dynamic analysis of cerebral blood flow regulation in humans",
            "citation": "IEEE Transactions on Biomedical Engineering. 2000;47(3):419-423. doi:10.1109/10.827312",
            "url": "https://stacks.cdc.gov/view/cdc/198526/cdc_198526_DS1.pdf",
        },
        {
            "authors": "Peng T, Rowley AB, Ainslie PN, Poulin MJ, Payne SJ",
            "title": "Multivariate system identification for cerebral autoregulation",
            "citation": "Annals of Biomedical Engineering. 2008;36(2):308-320. doi:10.1007/s10439-007-9412-9",
            "url": "https://ora.ox.ac.uk/objects/uuid:20fad100-eded-4fc7-acbc-decd6625fcbb",
        },
    ],
    "mci": [
        {
            "authors": "Kisler K, Nelson AR, Montagne A, Zlokovic BV",
            "title": "Cerebral blood flow regulation and neurovascular dysfunction in Alzheimer disease",
            "citation": "Nature Reviews Neuroscience. 2017;18(7):419-434. doi:10.1038/nrn.2017.48",
            "url": "https://doi.org/10.1038/nrn.2017.48",
        },
        {
            "authors": "Beishon L, Haunton VJ, Panerai RB, Robinson TG",
            "title": "Cerebral hemodynamics in mild cognitive impairment: A systematic review",
            "citation": "Journal of Alzheimer's Disease. 2017;59(1):369-385. doi:10.3233/JAD-170181",
            "url": "https://doi.org/10.3233/JAD-170181",
        },
        {
            "authors": "Gommer ED, Martens EGHJ, Aalten P, et al.",
            "title": "Dynamic cerebral autoregulation in subjects with Alzheimer's disease, mild cognitive impairment, and controls: Evidence for increased peripheral vascular resistance with possible predictive value",
            "citation": "Journal of Alzheimer's Disease. 2012;30(4):805-813. doi:10.3233/JAD-2012-111628",
            "url": "https://pubmed.ncbi.nlm.nih.gov/22460326/",
        },
        {
            "authors": "Tarumi T, Dunsky DI, Khan MA, et al.",
            "title": "Dynamic cerebral autoregulation and tissue oxygenation in amnestic mild cognitive impairment",
            "citation": "Journal of Alzheimer's Disease. 2014;41(3):765-778. doi:10.3233/JAD-132018",
            "url": "https://doi.org/10.3233/JAD-132018",
        },
        {
            "authors": "Heutz R, Claassen J, Feiner S, et al.",
            "title": "Dynamic cerebral autoregulation in Alzheimer's disease and mild cognitive impairment: A systematic review",
            "citation": "Journal of Cerebral Blood Flow & Metabolism. 2023;43(8):1223-1236. doi:10.1177/0271678X231173449",
            "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC10369144/",
        },
        {
            "authors": "Marmarelis VZ, Shin DC, Orme ME, Zhang R",
            "title": "Model-based physiomarkers of cerebral hemodynamics in patients with mild cognitive impairment",
            "citation": "Medical Engineering & Physics. 2014;36(5):628-637. doi:10.1016/j.medengphy.2014.02.025",
            "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC4076301/",
        },
        {
            "authors": "Marmarelis VZ, Shin DC, Tarumi T, Zhang R",
            "title": "Comparison of model-based indices of cerebral autoregulation and vasomotor reactivity using transcranial Doppler versus near-infrared spectroscopy in patients with amnestic mild cognitive impairment",
            "citation": "Journal of Alzheimer's Disease. 2017;56(1):89-105. doi:10.3233/JAD-161004",
            "url": "https://doi.org/10.3233/JAD-161004",
        },
        {
            "authors": "Hashem S, Yamashiro S, Joe E, Chui H, Marmarelis V",
            "title": "PRBS gas challenges reveal impaired chemoreflex and cholinergic dynamics in MCI",
            "citation": "Annals of Biomedical Engineering. 2026;54(8):2497-2507. doi:10.1007/s10439-026-04213-7",
            "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC13391744/",
        },
    ],
    "statistics": [
        {
            "authors": "Welch PD",
            "title": "The use of fast Fourier transform for the estimation of power spectra: A method based on time averaging over short, modified periodograms",
            "citation": "IEEE Transactions on Audio and Electroacoustics. 1967;15(2):70-73. doi:10.1109/TAU.1967.1161901",
            "url": "https://doi.org/10.1109/TAU.1967.1161901",
        },
        {
            "authors": "Mardia KV, Jupp PE",
            "title": "Directional Statistics",
            "citation": "Chichester, UK: Wiley; 2000. doi:10.1002/9780470316979",
            "url": "https://onlinelibrary.wiley.com/doi/book/10.1002/9780470316979",
        },
        {
            "authors": "Berens P",
            "title": "CircStat: A MATLAB toolbox for circular statistics",
            "citation": "Journal of Statistical Software. 2009;31(10):1-21. doi:10.18637/jss.v031.i10",
            "url": "https://www.jstatsoft.org/v31/i10",
        },
        {
            "authors": "Welch BL",
            "title": "The generalization of Student's problem when several different population variances are involved",
            "citation": "Biometrika. 1947;34(1-2):28-35. doi:10.1093/biomet/34.1-2.28",
            "url": "https://doi.org/10.1093/biomet/34.1-2.28",
        },
        {
            "authors": "Hedges LV",
            "title": "Distribution theory for Glass's estimator of effect size and related estimators",
            "citation": "Journal of Educational Statistics. 1981;6(2):107-128. doi:10.2307/1164588",
            "url": "https://doi.org/10.2307/1164588",
        },
        {
            "authors": "Benjamini Y, Hochberg Y",
            "title": "Controlling the false discovery rate: A practical and powerful approach to multiple testing",
            "citation": "Journal of the Royal Statistical Society Series B. 1995;57(1):289-300. doi:10.1111/j.2517-6161.1995.tb02031.x",
            "url": "https://doi.org/10.1111/j.2517-6161.1995.tb02031.x",
        },
    ],
    "cognitive": [
        {
            "authors": "Folstein MF, Folstein SE, McHugh PR",
            "title": "“Mini-mental state”: A practical method for grading the cognitive state of patients for the clinician",
            "citation": "Journal of Psychiatric Research. 1975;12(3):189-198. doi:10.1016/0022-3956(75)90026-6",
            "url": "https://pubmed.ncbi.nlm.nih.gov/1202204/",
        },
        {
            "authors": "Nasreddine ZS, Phillips NA, Bédirian V, et al.",
            "title": "The Montreal Cognitive Assessment, MoCA: A brief screening tool for mild cognitive impairment",
            "citation": "Journal of the American Geriatrics Society. 2005;53(4):695-699. doi:10.1111/j.1532-5415.2005.53221.x",
            "url": "https://pubmed.ncbi.nlm.nih.gov/15817019/",
        },
    ],
}


document = Document(DOCUMENT_PATH)
paragraphs = document.paragraphs
body_template = paragraphs[4]

# Replace only the temporary numbered citations. The user's revised wording stays intact.
replace_with_linked_parts(
    paragraphs[4],
    [
        ("The brain requires continuous blood flow even as arterial pressure changes from beat to beat. Cerebral autoregulation describes the physiological processes that limit how strongly these pressure changes alter cerebral perfusion (", None),
        ("Regulation of Cerebral Blood Flow in Humans", "https://pmc.ncbi.nlm.nih.gov/articles/PMC8576366/"),
        ("). Dynamic cerebral autoregulation focuses on responses that occur over seconds to minutes. TFA is commonly used to study these responses by treating MAP as the input and CBFV as the output (", None),
        ("Transfer Function Analysis of Dynamic Cerebral Autoregulation", "https://doi.org/10.1177/0271678X221119760"),
        ("). The resulting transfer function is complex-valued. Gain describes how strongly changes in MAP appear in CBFV, while phase describes their timing relationship at each frequency. Coherence is also used to describe how consistently the two signals are linearly related.", None),
    ],
)

replace_with_linked_parts(
    paragraphs[5],
    [
        ("A MAP-only model is useful, but MAP is not the only input that can influence CBFV. Carbon dioxide, being a potent vasodilator, has a strong effect on cerebrovascular tone, and PETCO₂ provides a practical continuous estimate of carbon dioxide during resting recordings (", None),
        ("Regulation of the Cerebral Circulation by Arterial Carbon Dioxide", "https://pubmed.ncbi.nlm.nih.gov/31187899/"),
        ("). When MAP and PETCO₂ share frequency content, a MAP-only SISO estimate can include part of the PETCO₂-related response. The same issue applies to a PETCO₂-only SISO estimate. Therefore, two separate SISO models describe pairwise relationships, but they do not separate the contribution of one measured input from the other.", None),
    ],
)

replace_with_linked_parts(
    paragraphs[6],
    [
        ("Previous studies have used multivariable models to evaluate MAP and carbon dioxide together (", None),
        ("Multivariate Dynamic Analysis of Cerebral Blood Flow Regulation", "https://stacks.cdc.gov/view/cdc/198526/cdc_198526_DS1.pdf"),
        ("; ", None),
        ("Multivariate System Identification for Cerebral Autoregulation", "https://ora.ox.ac.uk/objects/uuid:20fad100-eded-4fc7-acbc-decd6625fcbb"),
        ("). These studies provide a strong basis for using a two-input model. However, the use of MISO TFA is not yet organized into one clear and reproducible workflow for resting cerebral hemodynamic recordings. Studies may use different preprocessing choices, spectral settings, model solutions, diagnostic criteria, and reporting methods. In addition, higher multiple coherence shows that the model fits the measured output more closely, but it does not show that the separated MAP and PETCO₂ pathways are more accurate.", None),
    ],
)

# Shorten the overview without changing the user's preceding introduction.
overview_text = (
    "This study has three connected parts. First, the MISO TFA pipeline is developed and "
    "documented from signal preprocessing through spectral estimation, the joint MISO "
    "solution, model diagnostics, statistical testing, and organized outputs. Second, "
    "known-truth simulations are used to determine when MISO or SISO more accurately "
    "recovers the MAP and PETCO₂ pathways as the physiological system and observation "
    "conditions change. Third, NC recordings are placed within the simulated operating "
    "space to determine whether their duration, input relationships, excitation, and "
    "conditioning support interpretable MISO estimates."
)
overview_metrics = (
    "Model performance is summarized using gain error, wrapped phase error, normalized "
    "complex error, and model advantage. Positive model advantage favors MISO, while "
    "negative model advantage favors SISO. Statistical significance will always be "
    "interpreted together with this direction and the size of the difference."
)

overview_run = paragraphs[19].runs[0] if paragraphs[19].runs else body_template.runs[0]
clear_paragraph(paragraphs[19])
run = paragraphs[19].add_run(overview_text)
copy_run_style(overview_run, run)

clear_paragraph(paragraphs[20])
run = paragraphs[20].add_run(overview_metrics)
copy_run_style(body_template.runs[0], run)
run.bold = False
run.italic = False

# Remove the longer duplicated overview, including its equation and NC paragraph.
for paragraph in list(document.paragraphs[21:25]):
    delete_paragraph(paragraph)

# Replace the former selected reference list with a complete working library.
reference_heading = next(
    paragraph for paragraph in document.paragraphs if paragraph.text == "Selected References"
)
set_plain_heading(
    reference_heading,
    "Reference Papers Organized by Purpose",
    body_template.runs[0],
    size=16,
)

reference_heading_index = next(
    index
    for index, paragraph in enumerate(document.paragraphs)
    if paragraph._p is reference_heading._p
)
for paragraph in list(document.paragraphs[reference_heading_index + 1 :]):
    delete_paragraph(paragraph)

intro = document.add_paragraph()
intro_run = intro.add_run(
    "This working library includes every source cited in either the current approach-paper "
    "draft or the earlier combined MISO and MCI draft. Each paper title is linked to an "
    "online paper record or full-text page so that the source can be opened and downloaded."
)
copy_run_style(body_template.runs[0], intro_run)

add_reference_group(
    document,
    body_template,
    "Cerebral Blood Flow, dCA, and TFA Foundations",
    "These sources establish the physiology of cerebral blood flow regulation, the accepted definition and reporting of TFA, related methodological recommendations, and the basis for measuring CBFV with transcranial Doppler ultrasound.",
    references["foundations"],
)
add_reference_group(
    document,
    body_template,
    "Carbon Dioxide and Multivariable Cerebral Hemodynamic Modeling",
    "These papers support PETCO₂ as a meaningful cerebrovascular input and provide the direct basis for estimating MAP, PETCO₂, and CBFV together rather than as isolated signal pairs.",
    references["multivariable"],
)
add_reference_group(
    document,
    body_template,
    "MCI, Alzheimer's Disease, and Biological-Marker Background",
    "These sources mainly support the earlier combined draft and the separate biological-marker paper. They describe cerebral hemodynamic findings in MCI and Alzheimer's disease and provide examples of model-based physiomarkers and carbon-dioxide-sensitive approaches.",
    references["mci"],
)
add_reference_group(
    document,
    body_template,
    "Spectral Estimation and Statistical Methods",
    "These sources support Welch spectral estimation, circular treatment of phase, unequal-variance testing, standardized effect size, and false-discovery-rate control across repeated frequency-bin comparisons.",
    references["statistics"],
)
add_reference_group(
    document,
    body_template,
    "Cognitive Assessment Methods from the Earlier Combined Draft",
    "These papers define the MMSE and MoCA measures used to characterize cognitive status in the earlier clinical draft. They are not central to the approach-paper introduction but remain part of the complete working reference library.",
    references["cognitive"],
)

document.save(DOCUMENT_PATH)
print("Revised study overview and reference library")
