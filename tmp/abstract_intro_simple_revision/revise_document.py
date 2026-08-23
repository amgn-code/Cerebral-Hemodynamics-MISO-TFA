from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCUMENT_PATH = Path(__file__).with_name("Abstract & Introduction.docx")


def find_paragraph(document, beginning):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(beginning)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning with {beginning!r}; "
            f"found {len(matches)}."
        )
    return matches[0]


def replace_plain_paragraph(paragraph, text):
    """Replace body text while preserving the paragraph formatting."""
    paragraph.clear()
    paragraph.add_run(text)


def set_run_text(run_element, text):
    """Replace run text while retaining any existing run formatting."""
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)
    text_element = run_element.makeelement(qn("w:t"))
    text_element.text = text
    run_element.append(text_element)


def replace_around_only_hyperlink(paragraph, before_text, after_text):
    """Replace surrounding prose while retaining the linked paper title."""
    hyperlinks = [
        child
        for child in paragraph._p
        if child.tag == qn("w:hyperlink")
    ]
    if len(hyperlinks) != 1:
        raise RuntimeError(
            "Expected exactly one hyperlink in paragraph beginning with "
            f"{paragraph.text[:50]!r}."
        )
    hyperlink = hyperlinks[0]
    children = list(paragraph._p)
    hyperlink_index = children.index(hyperlink)
    runs_before = [
        child
        for child in children[:hyperlink_index]
        if child.tag == qn("w:r")
    ]
    runs_after = [
        child
        for child in children[hyperlink_index + 1 :]
        if child.tag == qn("w:r")
    ]
    if not runs_before or not runs_after:
        raise RuntimeError("Expected text runs before and after the hyperlink.")

    set_run_text(runs_before[0], before_text)
    for run in runs_before[1:]:
        paragraph._p.remove(run)

    set_run_text(runs_after[0], after_text)
    for run in runs_after[1:]:
        paragraph._p.remove(run)


document = Document(DOCUMENT_PATH)
hyperlinks_before = len(
    document.element.body.findall(".//w:hyperlink", document.element.nsmap)
)

abstract_main = find_paragraph(document, "Transfer Function Analysis (TFA)")
replace_plain_paragraph(
    abstract_main,
    "Transfer Function Analysis (TFA) of dynamic cerebral autoregulation "
    "(dCA) commonly treats mean arterial pressure (MAP) as the single input "
    "affecting cerebral blood flow velocity (CBFV). This single-input-single-"
    "output (SISO) model describes the gain and phase relationship between MAP "
    "and CBFV. However, end-tidal carbon dioxide (PETCO₂) can also influence "
    "CBFV. A multiple-input-single-output (MISO) model can estimate the "
    "MAP-to-CBFV and PETCO₂-to-CBFV pathways together while accounting for "
    "the relationship between the two inputs. Previous work suggests that "
    "multiple-input models may explain more of the measured CBFV signal, but "
    "better model fit alone does not show whether the two pathways were "
    "estimated correctly. In this study, we develop a reproducible MISO TFA "
    "pipeline for preprocessing, spectral estimation, the joint MISO solution, "
    "gain and phase calculation, model diagnostics, statistical comparison, "
    "and reporting. The pipeline will be evaluated using a known-truth "
    "simulation. For each simulated family, the true MAP-to-CBFV and "
    "PETCO₂-to-CBFV transfer functions are defined first. Generated MAP and "
    "PETCO₂ signals are then passed through these pathways and combined to "
    "create CBFV. MISO and SISO models are estimated using only the generated "
    "MAP, PETCO₂, and CBFV signals, and the estimated transfer functions are "
    "compared directly with the transfer functions used to generate CBFV. "
    "Recording duration, noise, timing alignment, input relationships, and "
    "estimator settings are varied without changing the true pathways. "
    "Accuracy will be evaluated using gain error, wrapped phase error, "
    "normalized complex pathway error, and model advantage.",
)

abstract_end = find_paragraph(
    document,
    "Transfer functions estimated from one simulated recording",
)
replace_plain_paragraph(
    abstract_end,
    "Properties of resting baseline recordings from cognitively normal (NC) "
    "adults will then be compared with the simulation results. These properties "
    "include recording duration, the relationship between MAP and PETCO₂, "
    "PETCO₂ fluctuation size, and the numerical stability of the MISO "
    "solution. This comparison will show whether the recordings resemble "
    "simulated conditions in which the two pathways could be separated "
    "reliably. The purpose of this paper is to develop and explain a "
    "reproducible MISO TFA pipeline, identify when MISO improves pathway "
    "recovery, determine when SISO remains sufficient, and recognize when the "
    "available data may not support reliable pathway separation.",
)

panerai = find_paragraph(
    document,
    "Previous studies have already demonstrated",
)
replace_around_only_hyperlink(
    panerai,
    "Previous studies have shown the potential value of modeling MAP and "
    "PETCO₂ together. Panerai et al. used multivariate finite impulse "
    "response filters to estimate the contributions of beat-to-beat MAP and "
    "breath-to-breath PETCO₂ fluctuations to resting CBFV variability (",
    "). MAP explained a substantial portion of the measured CBFV variability, "
    "and adding PETCO₂ reduced the model mean squared error. These results "
    "suggest that spontaneous PETCO₂ fluctuations can help explain resting "
    "CBFV variability that is not captured by pressure alone.",
)

prior_limits = find_paragraph(
    document,
    "However, lower CBFV prediction error",
)
replace_around_only_hyperlink(
    prior_limits,
    "However, lower CBFV prediction error does not necessarily mean that the "
    "separate MAP and PETCO₂ transfer functions were estimated correctly. "
    "MAP, PETCO₂, and CBFV can be measured in human recordings, but the true "
    "MAP-to-CBFV and PETCO₂-to-CBFV pathways are not independently known. "
    "Panerai et al. could therefore test whether PETCO₂ improved the "
    "description of total CBFV, but they could not directly compare the "
    "estimated pathways with known physiological pathways. Peng et al. also "
    "examined multivariable system identification for cerebral autoregulation (",
    "). Together, these studies provide a basis for two-input modeling, but "
    "they do not fully answer when a frequency-domain MISO model can separate "
    "the two pathways accurately.",
)

workflow_need = find_paragraph(document, "The use of MISO TFA is not yet")
replace_plain_paragraph(
    workflow_need,
    "A clear and consistent MISO TFA workflow is needed for this study because "
    "preprocessing, spectral settings, the model solution, and diagnostic "
    "criteria can all affect the results. Higher multiple coherence indicates "
    "that a model explains the measured output more closely, but it does not "
    "show by itself that the MAP and PETCO₂ pathways were separated "
    "accurately. The pipeline must therefore document these choices and test "
    "the estimated pathways against known truth.",
)

pipeline = find_paragraph(document, "The first goal of this paper")
replace_plain_paragraph(
    pipeline,
    "The first goal of this paper is therefore to develop and document the "
    "MISO TFA pipeline. The pipeline begins with MAP, PETCO₂, and CBFV "
    "signals and applies a consistent preprocessing and spectral estimation "
    "procedure. The MAP and PETCO₂ auto-spectra, their cross-spectrum, and "
    "their cross-spectra with CBFV are then used in one joint MISO solution. "
    "The pipeline returns MAP and PETCO₂ gain and phase, coherence measures, "
    "input power, and diagnostics showing whether the two inputs can be "
    "separated reliably. It also creates the statistical comparisons, figures, "
    "tables, and source data needed to understand and reproduce the analysis. "
    "Each major choice will be explained so that the workflow can be inspected "
    "and modified.",
)

figure_two_note = find_paragraph(document, "Figure note: Show the flow from")
replace_plain_paragraph(
    figure_two_note,
    "Figure note: Show the flow from MAP, PETCO₂, and CBFV signals through "
    "preprocessing, Welch spectral estimation, the joint MISO solution, gain "
    "and phase calculation, checks of input separation and numerical stability, "
    "statistical analysis, and organized figures and source-data exports.",
)

second_input = find_paragraph(document, "Adding a second input")
replace_plain_paragraph(
    second_input,
    "Adding a second input does not automatically produce a more accurate "
    "model. MISO must estimate two pathway coefficients at each frequency from "
    "the same finite recording. If MAP and PETCO₂ become too similar within "
    "a frequency region, the model can have difficulty separating their "
    "effects. In this situation, small spectral errors can produce large "
    "changes in the estimated transfer functions. Short recordings, weak "
    "PETCO₂ fluctuations, measurement noise, signal misalignment, and "
    "estimator settings may also affect this balance. A useful pipeline must "
    "therefore report the recording conditions and numerical diagnostics that "
    "support each estimate.",
)

simulation = find_paragraph(
    document,
    "Subject recordings cannot show the true MAP-to-CBFV",
)
replace_plain_paragraph(
    simulation,
    "Subject recordings cannot show the true MAP-to-CBFV or PETCO₂-to-CBFV "
    "transfer functions. If MISO and SISO produce different results in an "
    "observed recording, that difference alone cannot show which model is "
    "closer to the underlying physiology. A known-truth simulation makes this "
    "comparison possible. For each simulated family, the true MAP and PETCO₂ "
    "transfer functions are defined first. MAP and PETCO₂ signals are then "
    "generated, passed through the true pathways, and combined to create CBFV. "
    "MISO and SISO models are estimated using only these generated signals. "
    "Their estimated gain, phase, and complex transfer functions can then be "
    "compared directly with the true transfer functions used to generate CBFV. "
    "Recording duration, noise, input relationships, timing, and estimator "
    "settings are varied to determine how each condition affects pathway "
    "recovery.",
)

objectives = find_paragraph(
    document,
    "This paper therefore has three connected objectives.",
)
replace_plain_paragraph(
    objectives,
    "This paper therefore has three connected objectives. The first is to "
    "develop and explain the MISO TFA pipeline and its main choices. The second "
    "is to use known-truth simulations to determine when MISO or SISO more "
    "accurately recovers the MAP and PETCO₂ pathways. The third is to compare "
    "measurable properties of the NC recordings with the simulation results and "
    "determine whether the available recordings resemble conditions that "
    "support reliable pathway separation. The overall goal is to provide a "
    "clear and reproducible workflow for future studies.",
)

hyperlinks_after = len(
    document.element.body.findall(".//w:hyperlink", document.element.nsmap)
)
if hyperlinks_after != hyperlinks_before:
    raise RuntimeError(
        f"Hyperlink count changed from {hyperlinks_before} to "
        f"{hyperlinks_after}."
    )

document.save(DOCUMENT_PATH)
print(f"Updated {DOCUMENT_PATH}")
print(f"Hyperlinks preserved: {hyperlinks_after}")
