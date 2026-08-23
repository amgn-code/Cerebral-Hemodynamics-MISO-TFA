from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph


DOCUMENT = Path(__file__).with_name("Abstract & Introduction.docx")


ABSTRACT = (
    "Transfer Function Analysis (TFA) of dynamic cerebral autoregulation (dCA) "
    "generally treats mean arterial pressure (MAP) as a single input affecting "
    "the output of cerebral blood flow velocity (CBFV). This single-input-single-"
    "output (SISO) model provides the complex relationship, including gain and "
    "phase, between MAP and CBFV. However, end-tidal carbon dioxide (PETCO₂) can "
    "also influence CBFV. A multiple-input-single-output (MISO) model can jointly "
    "estimate the MAP-to-CBFV and PETCO₂-to-CBFV transfer functions while "
    "including the cross-spectral relationship between the two inputs. Previous "
    "studies have shown higher multiple coherence and lower CBFV prediction error "
    "with multiple-input models, but these measures do not show whether the "
    "separate pathways are more accurate or instead reflect the greater flexibility "
    "of a model with an additional input. In this study, we develop and validate a "
    "transparent MISO TFA pipeline that provides a consistent procedure for "
    "preprocessing, spectral estimation, the joint MISO solution, gain and phase "
    "calculation, model diagnostics, statistical comparison, and reporting. Short "
    "recordings, weak PETCO₂ fluctuations, correlated inputs, measurement noise, "
    "timing errors, and poorly conditioned spectral matrices may offset the "
    "conceptual benefit of separating the two pathways. The pipeline is therefore "
    "evaluated using a known-truth simulation framework. Each simulated family "
    "represents a fixed physiological system, while recording duration, noise, "
    "timing alignment, and estimator settings are varied without changing the "
    "underlying truth. Accuracy will be evaluated using frequency-resolved gain "
    "error, wrapped phase error, normalized complex pathway error, and model "
    "advantage. Transfer functions estimated from one simulated realization will "
    "also be applied to an independent realization to calculate normalized CBFV "
    "prediction error and determine whether improved prediction is accompanied by "
    "more accurate pathway recovery. Resting baseline recordings from cognitively "
    "normal (NC) adults will then be placed within the simulated operating space to "
    "determine whether their duration, input relationships, excitation, and "
    "conditioning support interpretable MISO estimates. The purpose of this paper "
    "is to develop a reproducible MISO TFA pipeline, explain the reasoning behind "
    "its main choices, and establish when MISO is justified, when SISO remains "
    "sufficient, and when neither "
    "separated pathway should be interpreted without additional data."
)


def clear_paragraph(paragraph):
    """Remove text and hyperlinks while preserving paragraph formatting."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def copy_run_style(source_run, target_run):
    if source_run is not None and source_run._element.rPr is not None:
        target_run._element.insert(0, deepcopy(source_run._element.rPr))


def replace_text(paragraph, text):
    source_run = paragraph.runs[0] if paragraph.runs else None
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    copy_run_style(source_run, run)


def add_hyperlink(paragraph, text, url, source_run=None):
    relationship_id = paragraph.part.relate_to(
        url, RT.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    if source_run is not None and source_run._element.rPr is not None:
        run_properties = deepcopy(source_run._element.rPr)
    else:
        run_properties = OxmlElement("w:rPr")

    color = run_properties.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        run_properties.append(color)
    color.set(qn("w:val"), "0563C1")

    underline = run_properties.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        run_properties.append(underline)
    underline.set(qn("w:val"), "single")

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_with_linked_parts(paragraph, parts):
    source_run = paragraph.runs[0] if paragraph.runs else None
    source_properties = (
        deepcopy(source_run._element.rPr)
        if source_run is not None and source_run._element.rPr is not None
        else None
    )
    clear_paragraph(paragraph)

    for text, url in parts:
        if url:
            add_hyperlink(paragraph, text, url, source_run)
        else:
            run = paragraph.add_run(text)
            if source_properties is not None:
                run._element.insert(0, deepcopy(source_properties))


def insert_paragraph_after(paragraph, text=""):
    """Insert a paragraph with the same paragraph and run formatting."""
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if paragraph._p.pPr is not None:
        new_paragraph._p.insert(0, deepcopy(paragraph._p.pPr))
    if text:
        run = new_paragraph.add_run(text)
        source_run = paragraph.runs[0] if paragraph.runs else None
        copy_run_style(source_run, run)
    return new_paragraph


def copy_paragraph_format(source, target):
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def insert_linked_paragraph_after(paragraph, parts):
    new_paragraph = insert_paragraph_after(paragraph)
    source_run = paragraph.runs[0] if paragraph.runs else None
    for text, url in parts:
        if url:
            add_hyperlink(new_paragraph, text, url, source_run)
        else:
            run = new_paragraph.add_run(text)
            copy_run_style(source_run, run)
    return new_paragraph


def delete_paragraph(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def paragraph_with_start(document, start):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(start):
            return paragraph
    raise ValueError(f"Paragraph not found: {start}")


def paragraph_with_text(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def main():
    document = Document(DOCUMENT)

    # Keep the user's overall abstract structure but remove repetition and add
    # the distinction between output prediction and pathway recovery.
    replace_text(document.paragraphs[2], ABSTRACT)

    carbon_paragraph = paragraph_with_start(document, "A MAP-only model is useful")
    replace_with_linked_parts(
        carbon_paragraph,
        [
            (
                "A MAP-only model is useful, but MAP is not the only input that can "
                "influence CBFV. Carbon dioxide is a potent vasodilator with a strong "
                "effect on cerebrovascular tone, and PETCO₂ provides a practical "
                "breath-to-breath surrogate for changes in arterial carbon dioxide "
                "during resting recordings (",
                None,
            ),
            (
                "Regulation of the Cerebral Circulation by Arterial Carbon Dioxide",
                "https://pubmed.ncbi.nlm.nih.gov/31187899/",
            ),
            (
                "). When MAP and PETCO₂ share frequency content, a MAP-only SISO "
                "estimate can include part of the PETCO₂-related response. The same "
                "issue applies to a PETCO₂-only SISO estimate. Therefore, two "
                "separate SISO models describe pairwise relationships, but they do "
                "not separate the contribution of one measured input from the other.",
                None,
            ),
        ],
    )

    prior_work = paragraph_with_start(document, "Previous studies have used")
    replace_with_linked_parts(
        prior_work,
        [
            (
                "Previous studies have already demonstrated the potential value of "
                "modeling MAP and PETCO₂ together. Panerai et al. used multivariate "
                "finite impulse response filters to estimate the contributions of "
                "beat-to-beat MAP and breath-to-breath PETCO₂ fluctuations to "
                "resting CBFV variability (",
                None,
            ),
            (
                "Multivariate Dynamic Analysis of Cerebral Blood Flow Regulation",
                "https://doi.org/10.1109/10.827312",
            ),
            (
                "). MAP explained a substantial portion of the observed variability, "
                "while adding PETCO₂ significantly reduced the model mean squared "
                "error and produced a physiologically reasonable carbon dioxide "
                "response. This work established that spontaneous PETCO₂ "
                "fluctuations may explain CBFV variability that is not captured by a "
                "pressure-only model.",
                None,
            ),
        ],
    )

    limitation = insert_linked_paragraph_after(
        prior_work,
        [
            (
                "However, a reduction in CBFV prediction error does not necessarily "
                "show that the separate MAP and PETCO₂ pathways have been recovered "
                "accurately. Because the true pathways were unknown in the human "
                "recordings, the study could not determine whether the two-input model "
                "correctly separated the two effects or simply described the total "
                "CBFV output more closely. The MAP and PETCO₂ inputs were also not "
                "significantly cross-correlated, leaving uncertainty about pathway "
                "recovery when the inputs share substantial frequency content or the "
                "input spectral matrix becomes poorly conditioned. Peng et al. further "
                "supported the use of multivariable system identification for cerebral "
                "autoregulation (",
                None,
            ),
            (
                "Multivariate System Identification for Cerebral Autoregulation",
                "https://ora.ox.ac.uk/objects/uuid:20fad100-eded-4fc7-acbc-decd6625fcbb",
            ),
            (
                "), but a clear and reproducible frequency-domain workflow with "
                "known-truth pathway validation is still needed.",
                None,
            ),
        ],
    )

    workflow_gap = insert_paragraph_after(
        limitation,
        "The use of MISO TFA is not yet organized into one clear and reproducible "
        "workflow for resting cerebral hemodynamic recordings. Studies may use "
        "different preprocessing choices, spectral settings, model solutions, "
        "diagnostic criteria, and reporting methods. In addition, higher multiple "
        "coherence shows that a model explains the measured output more closely, but "
        "it does not show that the separated MAP and PETCO₂ pathways are more "
        "accurate.",
    )
    workflow_gap.paragraph_format.keep_with_next = False

    conditioning = paragraph_with_start(document, "Adding a second input")
    replace_text(
        conditioning,
        "Adding a second input does not automatically produce a more accurate model. "
        "MISO must estimate two pathway coefficients at each frequency from the same "
        "finite recording. If MAP and PETCO₂ become too similar within a frequency "
        "region, the input spectral matrix can become poorly conditioned and small "
        "spectral errors can produce large changes in the separated transfer "
        "functions. Short recordings, weak PETCO₂ fluctuations, measurement noise, "
        "signal misalignment, and estimator settings may also affect this balance. A "
        "useful pipeline must therefore report the conditions that support the MISO "
        "estimates, not only the estimates themselves.",
    )

    truth_paragraph = paragraph_with_start(document, "Subject recordings cannot show")
    replace_text(
        truth_paragraph,
        "Subject recordings cannot show the true MAP-to-CBFV or PETCO₂-to-CBFV "
        "transfer functions. If MISO and SISO produce different results in an "
        "observed recording, that difference alone cannot show which model is closer "
        "to the underlying physiology. A known-truth simulation makes this comparison "
        "possible. MAP and PETCO₂ inputs can be generated with assigned "
        "relationships, passed through predefined transfer functions, and combined "
        "into a CBFV output. Transfer functions estimated from one realization can be "
        "compared directly with the known pathways and then applied to an independent "
        "realization from the same physiological family. This allows pathway accuracy "
        "and normalized CBFV prediction error to be evaluated together while recording "
        "duration, noise, input relationships, timing, and estimator settings are "
        "varied.",
    )

    overview = paragraph_with_start(document, "This study has three connected parts")
    replace_text(
        overview,
        "This study has three connected parts. First, the MISO TFA pipeline is "
        "developed and documented from signal preprocessing through spectral "
        "estimation, the joint MISO solution, model diagnostics, statistical testing, "
        "and organized outputs. Second, known-truth simulations are used to determine "
        "when MISO or SISO more accurately recovers the MAP and PETCO₂ pathways and "
        "whether improved CBFV prediction is accompanied by improved pathway recovery. "
        "Third, NC recordings are placed within the simulated operating space to "
        "determine whether their duration, input relationships, excitation, and "
        "conditioning support interpretable MISO estimates.",
    )

    metrics = paragraph_with_start(document, "Model performance is summarized")
    replace_text(
        metrics,
        "Model performance is summarized using gain error, wrapped phase error, "
        "normalized complex pathway error, normalized CBFV prediction error, and model "
        "advantage. Model advantage is calculated as the base-10 logarithm of the SISO "
        "error divided by the MISO error. Positive model advantage favors MISO, while "
        "negative model advantage favors SISO. Statistical significance will always "
        "be interpreted together with this direction and the size of the difference.",
    )

    foundation_purpose = paragraph_with_start(
        document, "Use in the papers: These sources establish the physiology"
    )
    replace_text(
        foundation_purpose,
        "Use in the papers: Claassen and the CARNet papers establish the physiology "
        "of cerebral blood flow regulation and the accepted definition and reporting "
        "of TFA. Aaslid provides the original basis for measuring CBFV with "
        "transcranial Doppler ultrasound and is relevant when the measurement method "
        "is described.",
    )
    if foundation_purpose.runs:
        foundation_purpose.runs[0].italic = True

    panerai_reference = paragraph_with_start(
        document,
        "Panerai RB, Simpson DM, Deverson ST, Mahony P, Hayes P, Evans DH.",
    )
    replace_with_linked_parts(
        panerai_reference,
        [
            (
                "Panerai RB, Simpson DM, Deverson ST, Mahony P, Hayes P, Evans DH. ",
                None,
            ),
            (
                "Multivariate dynamic analysis of cerebral blood flow regulation in humans",
                "https://doi.org/10.1109/10.827312",
            ),
            (
                ". IEEE Transactions on Biomedical Engineering. "
                "2000;47(3):419-423. doi:10.1109/10.827312",
                None,
            ),
        ],
    )

    statistics_heading = paragraph_with_text(
        document, "Spectral Estimation and Statistical Methods"
    )
    replace_text(
        statistics_heading,
        "Spectral Estimation and Paper 1 Statistical Methods",
    )
    if statistics_heading.runs:
        statistics_heading.runs[0].bold = True

    statistics_purpose = paragraph_with_start(
        document, "Use in the papers: These sources support Welch spectral estimation"
    )
    replace_text(
        statistics_purpose,
        "Use in the paper: These sources support Welch spectral estimation, the "
        "circular treatment of phase implemented directly from directional-statistics "
        "definitions, and false-discovery-rate control across related comparisons.",
    )
    if statistics_purpose.runs:
        statistics_purpose.runs[0].italic = True

    berens = paragraph_with_start(document, "Berens P. CircStat")
    welch_group = paragraph_with_start(document, "Welch BL. The generalization")
    hedges_group = paragraph_with_start(document, "Hedges LV. Distribution theory")
    bh_reference = paragraph_with_start(
        document, "Benjamini Y, Hochberg Y. Controlling"
    )
    delete_paragraph(berens)

    # Keep the Paper 1 sources together, then place the independent-group sources
    # in their own clearly labeled working-library section.
    mardia = paragraph_with_start(document, "Mardia KV, Jupp PE.")
    mardia._p.addnext(bh_reference._p)
    paper2_heading = insert_paragraph_after(
        bh_reference, "Independent Group Comparisons for the Biological-Marker Paper"
    )
    copy_paragraph_format(statistics_heading, paper2_heading)
    if paper2_heading.runs:
        paper2_heading.runs[0].bold = True
        paper2_heading.runs[0].italic = False
    paper2_heading.paragraph_format.space_before = statistics_heading.paragraph_format.space_before
    paper2_heading.paragraph_format.space_after = statistics_heading.paragraph_format.space_after

    paper2_purpose = insert_paragraph_after(
        paper2_heading,
        "Use in the paper: These sources support unequal-variance NC-MCI comparisons "
        "and standardized independent-group effect sizes in the separate biological-"
        "marker paper. They are not used in the primary Paper 1 simulation analysis.",
    )
    copy_paragraph_format(statistics_purpose, paper2_purpose)
    if paper2_purpose.runs:
        paper2_purpose.runs[0].bold = False
        paper2_purpose.runs[0].italic = True

    paper2_purpose._p.addnext(welch_group._p)
    welch_group._p.addnext(hedges_group._p)

    document.save(DOCUMENT)
    print("Revised abstract, introduction, study overview, and reference organization")


if __name__ == "__main__":
    main()
